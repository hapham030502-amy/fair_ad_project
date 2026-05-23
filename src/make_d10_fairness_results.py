from __future__ import annotations

from pathlib import Path
from typing import Dict, List, Tuple

import numpy as np
import pandas as pd
import matplotlib.pyplot as plt

from sklearn.ensemble import IsolationForest
from sklearn.metrics import average_precision_score, f1_score, roc_auc_score

from src.data_loader import load_processed_data


ROOT = Path(__file__).resolve().parents[1]

RESULTS_DIR = ROOT / "results"
D9_DIR = RESULTS_DIR / "d9"
D10_DIR = RESULTS_DIR / "d10"

ALL_RESULTS_FILE = RESULTS_DIR / "all_results.csv"
D9_COMPARISON_FILE = D9_DIR / "D9_baseline_vs_recommended_postprocessing.csv"
D9_PARETO_MEAN_FILE = D9_DIR / "d9_pareto_mean.csv"

SEEDS = [42, 123, 456, 789, 1011, 2026, 31415, 27182, 16180, 14142]

def safe_pr_auc(y_true: np.ndarray, scores: np.ndarray) -> float:
    y_true = np.asarray(y_true, dtype=int).ravel()
    scores = np.asarray(scores, dtype=float).ravel()

    if len(np.unique(y_true)) < 2:
        return 0.0

    return float(average_precision_score(y_true, scores))


def safe_roc_auc(y_true: np.ndarray, scores: np.ndarray) -> float:
    y_true = np.asarray(y_true, dtype=int).ravel()
    scores = np.asarray(scores, dtype=float).ravel()

    if len(np.unique(y_true)) < 2:
        return 0.5

    return float(roc_auc_score(y_true, scores))


def compute_fpr_fnr(
    y_true: np.ndarray,
    y_pred: np.ndarray,
    mask: np.ndarray,
) -> Tuple[float, float]:
    yt = y_true[mask]
    yp = y_pred[mask]

    fp = int(((yp == 1) & (yt == 0)).sum())
    tn = int(((yp == 0) & (yt == 0)).sum())
    fn = int(((yp == 0) & (yt == 1)).sum())
    tp = int(((yp == 1) & (yt == 1)).sum())

    fpr = fp / (fp + tn) if (fp + tn) > 0 else 0.0
    fnr = fn / (fn + tp) if (fn + tp) > 0 else 0.0

    return float(fpr), float(fnr)


def compute_fairness_metrics(
    y_true: np.ndarray,
    y_pred: np.ndarray,
    sensitive: np.ndarray,
) -> Dict[str, float]:
    y_true = np.asarray(y_true, dtype=int).ravel()
    y_pred = np.asarray(y_pred, dtype=int).ravel()
    sensitive = np.asarray(sensitive).ravel()

    groups = np.unique(sensitive)

    if len(groups) < 2:
        return {
            "delta_fpr": 0.0,
            "delta_fnr": 0.0,
            "eo_gap": 0.0,
        }

    fprs: List[float] = []
    fnrs: List[float] = []

    for g in groups:
        fpr, fnr = compute_fpr_fnr(y_true, y_pred, sensitive == g)
        fprs.append(fpr)
        fnrs.append(fnr)

    delta_fpr = float(max(fprs) - min(fprs))
    delta_fnr = float(max(fnrs) - min(fnrs))

    return {
        "delta_fpr": delta_fpr,
        "delta_fnr": delta_fnr,
        "eo_gap": delta_fpr + delta_fnr,
    }


def choose_threshold_by_f1(
    y_val: np.ndarray,
    val_scores: np.ndarray,
) -> Tuple[float, float]:
    """
    Chọn threshold trên validation set theo F1.
    Không dùng test set để chọn threshold.
    """
    y_val = np.asarray(y_val, dtype=int).ravel()
    val_scores = np.asarray(val_scores, dtype=float).ravel()

    percentiles = np.linspace(50, 99.5, 100)

    best_theta = float(np.percentile(val_scores, 95))
    best_f1 = -1.0

    for p in percentiles:
        theta = float(np.percentile(val_scores, p))
        pred = (val_scores >= theta).astype(int)
        f1 = float(f1_score(y_val, pred, zero_division=0))

        if f1 > best_f1:
            best_f1 = f1
            best_theta = theta

    return best_theta, best_f1


def compute_sample_weights_by_group(sensitive_train: np.ndarray) -> np.ndarray:
    """
    In-processing method: reweighting theo nhóm sensitive.

    Nhóm có ít mẫu hơn sẽ có trọng số cao hơn.
    Đây là lựa chọn đơn giản, phù hợp với yêu cầu D10:
    ít nhất có một in-processing method để so sánh với baseline và post-processing.
    """
    sensitive_train = np.asarray(sensitive_train).ravel()

    weights = np.ones(len(sensitive_train), dtype=float)
    groups = np.unique(sensitive_train)

    for g in groups:
        mask = sensitive_train == g
        n_g = int(mask.sum())

        if n_g > 0:
            weights[mask] = len(sensitive_train) / (len(groups) * n_g)

    weights = weights / np.mean(weights)

    return weights


def run_reweighted_isolation_forest(
    dataset: str,
    seed: int,
) -> Dict[str, float]:
    """
    In-processing: IsolationForest có sample_weight trong fit().
    Train trên X_train normal-only, chọn threshold trên validation,
    đánh giá cuối cùng trên test.
    """
    d = load_processed_data(dataset)

    X_train = d["X_train"]
    X_val = d["X_val"]
    X_test = d["X_test"]

    y_val = d["y_val"]
    y_test = d["y_test"]

    s_train = d["s_train"]
    s_test = d["s_test"]

    sample_weight = compute_sample_weights_by_group(s_train)

    model = IsolationForest(
        n_estimators=200,
        max_samples="auto",
        contamination="auto",
        random_state=seed,
        n_jobs=-1,
    )

    try:
        model.fit(X_train, sample_weight=sample_weight)
    except TypeError:
        print(
            "[WARN] Phiên bản scikit-learn không hỗ trợ sample_weight cho "
            "IsolationForest.fit(). Model sẽ fit không trọng số."
        )
        model.fit(X_train)

    # sklearn IsolationForest: điểm càng lớn càng bình thường.
    # Do đó lấy dấu âm để điểm càng lớn càng bất thường.
    val_scores = -model.decision_function(X_val)
    test_scores = -model.decision_function(X_test)

    theta, val_f1 = choose_threshold_by_f1(y_val, val_scores)
    y_pred = (test_scores >= theta).astype(int)

    out = {
        "dataset": dataset,
        "method": "In-processing",
        "method_detail": "Reweighted IsolationForest",
        "seed": seed,
        "threshold": theta,
        "val_f1": val_f1,
        "roc_auc": safe_roc_auc(y_test, test_scores),
        "pr_auc": safe_pr_auc(y_test, test_scores),
        "f1": float(f1_score(y_test, y_pred, zero_division=0)),
    }

    out.update(compute_fairness_metrics(y_test, y_pred, s_test))

    return out


def load_baseline_summary() -> pd.DataFrame:
    """
    Lấy baseline theo base model đã chọn ở D9.
    """
    if not ALL_RESULTS_FILE.exists():
        raise FileNotFoundError(
            f"Không tìm thấy {ALL_RESULTS_FILE}. Cần có kết quả D6 trước."
        )

    if not D9_COMPARISON_FILE.exists():
        raise FileNotFoundError(
            f"Không tìm thấy {D9_COMPARISON_FILE}. Cần chạy D9 trước."
        )

    all_results = pd.read_csv(ALL_RESULTS_FILE)
    d9_compare = pd.read_csv(D9_COMPARISON_FILE)

    rows = []

    for _, r in d9_compare.iterrows():
        dataset = r["dataset"]
        base_model = r["base_model"]

        sub = all_results[
            (all_results["dataset"] == dataset)
            & (all_results["model"] == base_model)
        ].copy()

        if sub.empty:
            raise ValueError(
                f"Không tìm thấy baseline cho dataset={dataset}, model={base_model}"
            )

        rows.append(
            {
                "dataset": dataset,
                "method": "Baseline",
                "method_detail": str(base_model),
                "roc_auc_mean": sub["roc_auc"].mean(),
                "roc_auc_std": sub["roc_auc"].std(),
                "pr_auc_mean": sub["pr_auc"].mean(),
                "pr_auc_std": sub["pr_auc"].std(),
                "f1_mean": sub["f1"].mean(),
                "f1_std": sub["f1"].std(),
                "delta_fpr_mean": sub["delta_fpr"].mean(),
                "delta_fpr_std": sub["delta_fpr"].std(),
                "delta_fnr_mean": sub["delta_fnr"].mean(),
                "delta_fnr_std": sub["delta_fnr"].std(),
                "eo_gap_mean": sub["eo_gap"].mean(),
                "eo_gap_std": sub["eo_gap"].std(),
            }
        )

    return pd.DataFrame(rows)


def load_postprocessing_summary() -> pd.DataFrame:
    """
    Lấy post-processing khuyến nghị từ D9.
    """
    if not D9_COMPARISON_FILE.exists():
        raise FileNotFoundError(
            f"Không tìm thấy {D9_COMPARISON_FILE}. Cần chạy D9 trước."
        )

    d9_compare = pd.read_csv(D9_COMPARISON_FILE)

    pareto = None

    if D9_PARETO_MEAN_FILE.exists():
        pareto = pd.read_csv(D9_PARETO_MEAN_FILE)

    rows = []

    for _, r in d9_compare.iterrows():
        dataset = r["dataset"]
        base_model = r["base_model"]
        rule = r["recommended_rule"]
        param = float(r["recommended_param"])

        delta_fpr = np.nan
        delta_fnr = np.nan
        delta_fpr_std = np.nan
        delta_fnr_std = np.nan
        pr_auc_std = np.nan
        f1_std = np.nan
        eo_gap_std = np.nan

        if pareto is not None:
            sub = pareto[
                (pareto["dataset"] == dataset)
                & (pareto["base_model"] == base_model)
                & (pareto["rule"] == rule)
            ].copy()

            if not sub.empty:
                sub["param_diff"] = (sub["param"] - param).abs()
                best = sub.sort_values("param_diff").iloc[0]

                delta_fpr = best.get("delta_fpr_mean", np.nan)
                delta_fnr = best.get("delta_fnr_mean", np.nan)
                delta_fpr_std = best.get("delta_fpr_std", np.nan)
                delta_fnr_std = best.get("delta_fnr_std", np.nan)
                pr_auc_std = best.get("pr_auc_std", np.nan)
                f1_std = best.get("f1_std", np.nan)
                eo_gap_std = best.get("eo_gap_std", np.nan)

        rows.append(
            {
                "dataset": dataset,
                "method": "Post-processing",
              "method_detail": f"{rule} - param={param:.4f}",
                "roc_auc_mean": np.nan,
                "roc_auc_std": np.nan,
                "pr_auc_mean": r["post_pr_auc"],
                "pr_auc_std": pr_auc_std,
                "f1_mean": r["post_f1"],
                "f1_std": f1_std,
                "delta_fpr_mean": delta_fpr,
                "delta_fpr_std": delta_fpr_std,
                "delta_fnr_mean": delta_fnr,
                "delta_fnr_std": delta_fnr_std,
                "eo_gap_mean": r["post_eo_gap"],
                "eo_gap_std": eo_gap_std,
            }
        )

    return pd.DataFrame(rows)


def run_inprocessing_summary() -> Tuple[pd.DataFrame, pd.DataFrame]:
    """
    Chạy in-processing reweighting theo 10 seed.
    """
    datasets = ["adult", "credit_default"]

    raw_rows = []

    for dataset in datasets:
        for seed in SEEDS:
            print(f"[D10] Running in-processing: dataset={dataset}, seed={seed}")
            raw_rows.append(
                run_reweighted_isolation_forest(dataset, seed)
            )

    raw = pd.DataFrame(raw_rows)

    summary = (
        raw
        .groupby(["dataset", "method", "method_detail"], as_index=False)
        .agg(
            roc_auc_mean=("roc_auc", "mean"),
            roc_auc_std=("roc_auc", "std"),
            pr_auc_mean=("pr_auc", "mean"),
            pr_auc_std=("pr_auc", "std"),
            f1_mean=("f1", "mean"),
            f1_std=("f1", "std"),
            delta_fpr_mean=("delta_fpr", "mean"),
            delta_fpr_std=("delta_fpr", "std"),
            delta_fnr_mean=("delta_fnr", "mean"),
            delta_fnr_std=("delta_fnr", "std"),
            eo_gap_mean=("eo_gap", "mean"),
            eo_gap_std=("eo_gap", "std"),
        )
    )

    return raw, summary


def add_tradeoff_analysis(summary: pd.DataFrame) -> pd.DataFrame:
    """
    Tính thay đổi so với baseline theo từng dataset.
    """
    df = summary.copy()

    df["pr_auc_change_vs_baseline"] = np.nan
    df["f1_change_vs_baseline"] = np.nan
    df["eo_gap_reduction_vs_baseline"] = np.nan

    for dataset, idx in df.groupby("dataset").groups.items():
        sub = df.loc[list(idx)]
        base = sub[sub["method"] == "Baseline"]

        if base.empty:
            continue

        base = base.iloc[0]

        df.loc[sub.index, "pr_auc_change_vs_baseline"] = (
            sub["pr_auc_mean"] - base["pr_auc_mean"]
        )
        df.loc[sub.index, "f1_change_vs_baseline"] = (
            sub["f1_mean"] - base["f1_mean"]
        )
        df.loc[sub.index, "eo_gap_reduction_vs_baseline"] = (
            base["eo_gap_mean"] - sub["eo_gap_mean"]
        )

    return df


def choose_best_method(summary: pd.DataFrame) -> pd.DataFrame:
    """
    Chọn method trade-off tốt nhất theo từng dataset.

    Nguyên tắc:
    - Ưu tiên giảm EO_gap.
    - Không để PR-AUC giảm quá mạnh.
    - Nếu PR-AUC gần tương đương, chọn EO_gap nhỏ hơn.
    """
    rows = []

    for dataset, sub in summary.groupby("dataset"):
        base = sub[sub["method"] == "Baseline"].iloc[0]

        tmp = sub.copy()
        tmp["passes_pr_guardrail"] = (
            tmp["pr_auc_mean"] >= base["pr_auc_mean"] - 0.01
        )

        candidates = tmp[tmp["passes_pr_guardrail"]].copy()

        if candidates.empty:
            candidates = tmp.copy()

        candidates = candidates.sort_values(
            ["eo_gap_mean", "pr_auc_mean"],
            ascending=[True, False],
        )

        best = candidates.iloc[0].copy()
        best["selection_reason"] = (
            "EO_gap thấp nhất trong nhóm không làm PR-AUC giảm quá 0.01 so với baseline"
        )

        rows.append(best)

    return pd.DataFrame(rows)


def plot_d10_comparison(summary: pd.DataFrame, out_path: Path) -> None:
    datasets = sorted(summary["dataset"].unique())
    methods = ["Baseline", "Post-processing", "In-processing"]

    x = np.arange(len(datasets))
    width = 0.24

    plt.figure(figsize=(9.5, 6.0))

    for i, method in enumerate(methods):
        vals = []

        for dataset in datasets:
            sub = summary[
                (summary["dataset"] == dataset)
                & (summary["method"] == method)
            ]

            if sub.empty:
                vals.append(np.nan)
            else:
                vals.append(float(sub.iloc[0]["eo_gap_mean"]))

        plt.bar(
            x + (i - 1) * width,
            vals,
            width,
            label=method,
        )

    plt.xticks(x, datasets)
    plt.ylabel("EO_gap - thấp hơn là tốt hơn")
    plt.title("D10 Comparison: Baseline vs Post-processing vs In-processing")
    plt.legend()
    plt.grid(axis="y", alpha=0.25)
    plt.tight_layout()
    plt.savefig(out_path, dpi=300, bbox_inches="tight")
    plt.close()


def format_float(x) -> str:
    if pd.isna(x):
        return ""
    return f"{float(x):.4f}"


def markdown_table(df: pd.DataFrame, cols: List[str]) -> str:
    d = df[cols].copy()

    for c in d.columns:
        if pd.api.types.is_float_dtype(d[c]):
            d[c] = d[c].map(format_float)

    header = "| " + " | ".join(d.columns) + " |"
    sep = "| " + " | ".join(["---"] * len(d.columns)) + " |"

    rows = []

    for _, row in d.iterrows():
        rows.append(
            "| " + " | ".join(str(row[c]) for c in d.columns) + " |"
        )

    return "\n".join([header, sep] + rows)


def write_report_md(
    summary: pd.DataFrame,
    best: pd.DataFrame,
    out_path: Path,
) -> None:
    lines: List[str] = []

    lines.append("# D10 - Fairness-aware Results")
    lines.append("")
    lines.append("## 1. Mục tiêu")
    lines.append(
        "D10 so sánh ba nhóm kết quả: Baseline, Post-processing và In-processing. "
        "Các chỉ số chính gồm PR-AUC, F1 và EO_gap."
    )
    lines.append("")
    lines.append("## 2. Thiết kế thực nghiệm")
    lines.append("- **Baseline**: lấy mô hình nền tốt nhất theo PR-AUC từ D6/D9.")
    lines.append("- **Post-processing**: lấy quy tắc hậu xử lý khuyến nghị từ D9.")
    lines.append(
        "- **In-processing**: dùng reweighted IsolationForest, trong đó mẫu ở nhóm sensitive "
        "ít hơn được gán trọng số cao hơn khi huấn luyện."
    )
    lines.append(
        "- Threshold của in-processing được chọn trên validation set theo F1; "
        "test set chỉ dùng để đánh giá cuối cùng."
    )
    lines.append("")
    lines.append("## 3. Bảng so sánh tổng hợp")
    lines.append(
        markdown_table(
            summary,
            [
                "dataset",
                "method",
                "method_detail",
                "pr_auc_mean",
                "f1_mean",
                "eo_gap_mean",
                "pr_auc_change_vs_baseline",
                "f1_change_vs_baseline",
                "eo_gap_reduction_vs_baseline",
            ],
        )
    )
    lines.append("")
    lines.append("## 4. Method trade-off tốt nhất theo từng dataset")
    lines.append(
        markdown_table(
            best,
            [
                "dataset",
                "method",
                "method_detail",
                "pr_auc_mean",
                "f1_mean",
                "eo_gap_mean",
                "pr_auc_change_vs_baseline",
                "eo_gap_reduction_vs_baseline",
                "selection_reason",
            ],
        )
    )
    lines.append("")
    lines.append("## 5. Nhận xét")
    for _, row in best.iterrows():
        lines.append(
            f"- **{row['dataset']}**: phương pháp có trade-off tốt nhất là "
            f"**{row['method']}** ({row['method_detail']}). "
            f"PR-AUC = {row['pr_auc_mean']:.4f}, F1 = {row['f1_mean']:.4f}, "
            f"EO_gap = {row['eo_gap_mean']:.4f}. "
            f"So với baseline, EO_gap thay đổi "
            f"{row['eo_gap_reduction_vs_baseline']:.4f}."
        )

    lines.append("")
    lines.append("## 6. Kết luận D10")
    lines.append(
        "D10 đã tạo được bảng so sánh Baseline, Post-processing và In-processing. "
        "Kết quả cho phép đánh giá trực tiếp trade-off giữa hiệu quả phát hiện bất thường "
        "và công bằng nhóm. Phần này là đầu vào cho D11, nơi tiến hành ablation study."
    )
    lines.append("")
    lines.append("## 7. File kết quả")
    lines.append("- `results/d10/D10_inprocessing_raw_results.csv`")
    lines.append("- `results/d10/D10_fairness_aware_results.csv`")
    lines.append("- `results/d10/D10_best_method_by_dataset.csv`")
    lines.append("- `results/d10/d10_eogap_comparison.png`")
    lines.append("- `results/d10/D10_Fairness_Aware_Results_Report.md`")
    lines.append("- `results/d10/D10_Fairness_Aware_Results_Report.docx` nếu đã cài `python-docx`")

    out_path.write_text("\n".join(lines), encoding="utf-8")


def try_write_docx(md_path: Path, docx_path: Path) -> None:
    try:
        from docx import Document
        from docx.shared import Pt
    except Exception:
        print(
            "[WARN] Chưa cài python-docx nên chỉ tạo file .md. "
            "Muốn có .docx thì chạy: pip install python-docx"
        )
        return

    doc = Document()
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(13)

    for line in md_path.read_text(encoding="utf-8").splitlines():
        line = line.strip()

        if not line:
            doc.add_paragraph("")
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.startswith("|"):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        else:
            doc.add_paragraph(line)

    doc.save(docx_path)


def check_inputs() -> None:
    if not ALL_RESULTS_FILE.exists():
        raise FileNotFoundError(
            f"Không tìm thấy {ALL_RESULTS_FILE}. Cần có D6 trước."
        )

    if not D9_COMPARISON_FILE.exists():
        raise FileNotFoundError(
            f"Không tìm thấy {D9_COMPARISON_FILE}. Cần chạy D9 trước."
        )


def main() -> None:
    check_inputs()

    D10_DIR.mkdir(parents=True, exist_ok=True)

    print("[D10] Loading baseline summary...")
    baseline = load_baseline_summary()

    print("[D10] Loading post-processing summary from D9...")
    post = load_postprocessing_summary()

    print("[D10] Running in-processing reweighting...")
    in_raw, in_summary = run_inprocessing_summary()

    in_raw.to_csv(
        D10_DIR / "D10_inprocessing_raw_results.csv",
        index=False,
    )

    summary = pd.concat(
        [baseline, post, in_summary],
        ignore_index=True,
        sort=False,
    )

    summary = add_tradeoff_analysis(summary)

    summary = summary.sort_values(
        ["dataset", "method"],
        ascending=[True, True],
    ).reset_index(drop=True)

    summary.to_csv(
        D10_DIR / "D10_fairness_aware_results.csv",
        index=False,
    )

    best = choose_best_method(summary)

    best.to_csv(
        D10_DIR / "D10_best_method_by_dataset.csv",
        index=False,
    )

    plot_d10_comparison(
        summary,
        D10_DIR / "d10_eogap_comparison.png",
    )

    report_md = D10_DIR / "D10_Fairness_Aware_Results_Report.md"

    write_report_md(summary, best, report_md)

    root_report_md = ROOT / "D10_Fairness_Aware_Results_Report.md"
    root_report_md.write_text(
        report_md.read_text(encoding="utf-8"),
        encoding="utf-8",
    )

    try_write_docx(
        report_md,
        D10_DIR / "D10_Fairness_Aware_Results_Report.docx",
    )

    try_write_docx(
        root_report_md,
        ROOT / "D10_Fairness_Aware_Results_Report.docx",
    )

    print("[OK] Đã tạo xong D10.")
    print(f"[OK] Kết quả nằm trong: {D10_DIR}")
    print("[OK] File chính:")
    print(" - results/d10/D10_inprocessing_raw_results.csv")
    print(" - results/d10/D10_fairness_aware_results.csv")
    print(" - results/d10/D10_best_method_by_dataset.csv")
    print(" - results/d10/d10_eogap_comparison.png")
    print(" - results/d10/D10_Fairness_Aware_Results_Report.md")
    print(" - results/d10/D10_Fairness_Aware_Results_Report.docx nếu có python-docx")


if __name__ == "__main__":
    main()
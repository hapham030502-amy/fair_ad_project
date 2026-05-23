from __future__ import annotations

import hashlib
import shutil
from pathlib import Path
from typing import Dict, List, Optional

import numpy as np
import pandas as pd


ROOT = Path(__file__).resolve().parents[1]

REPRO_DIR = ROOT / "reproducibility"
RESULTS_DIR = ROOT / "results"
DATA_PROCESSED_DIR = ROOT / "data" / "processed"
CONFIG_DIR = ROOT / "config"
SCRIPTS_DIR = ROOT / "scripts"

D9_DIR = RESULTS_DIR / "d9"
D10_DIR = RESULTS_DIR / "d10"
D11_DIR = RESULTS_DIR / "d11"

SEEDS = [42, 123, 456, 789, 1011, 2026, 31415, 27182, 16180, 14142]


def ensure_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def copy_if_exists(src: Path, dst: Path) -> bool:
    if not src.exists():
        return False

    ensure_dir(dst.parent)

    if src.is_dir():
        if dst.exists():
            shutil.rmtree(dst)
        shutil.copytree(src, dst)
    else:
        shutil.copy2(src, dst)

    return True


def sha256_file(path: Path) -> str:
    h = hashlib.sha256()

    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)

    return h.hexdigest()


def make_manifest() -> pd.DataFrame:
    rows: List[Dict[str, str]] = []

    for p in sorted(REPRO_DIR.rglob("*")):
        if not p.is_file():
            continue

        rel = p.relative_to(REPRO_DIR).as_posix()

        # Không đưa chính file manifest vào manifest để tránh vòng lặp hash.
        if rel == "MANIFEST_sha256.csv":
            continue

        rows.append(
            {
                "relative_path": rel,
                "size_bytes": str(p.stat().st_size),
                "sha256": sha256_file(p),
            }
        )

    df = pd.DataFrame(rows)
    df.to_csv(
        REPRO_DIR / "MANIFEST_sha256.csv",
        index=False,
        encoding="utf-8-sig",
    )
    return df


def copy_configs_and_scripts() -> None:
    ensure_dir(REPRO_DIR / "configs")
    ensure_dir(REPRO_DIR / "scripts")

    copy_if_exists(CONFIG_DIR / "config.yaml", REPRO_DIR / "configs" / "config.yaml")

    copy_if_exists(ROOT / "requirements.txt", REPRO_DIR / "requirements.txt")
    copy_if_exists(ROOT / "environment.yml", REPRO_DIR / "environment.yml")

    copy_if_exists(SCRIPTS_DIR / "run_all.sh", REPRO_DIR / "scripts" / "run_all.sh")
    copy_if_exists(SCRIPTS_DIR / "check_d1_d8.sh", REPRO_DIR / "scripts" / "check_d1_d8.sh")
    copy_if_exists(SCRIPTS_DIR / "run_d9.sh", REPRO_DIR / "scripts" / "run_d9.sh")
    copy_if_exists(SCRIPTS_DIR / "run_d10.sh", REPRO_DIR / "scripts" / "run_d10.sh")
    copy_if_exists(SCRIPTS_DIR / "run_d11.sh", REPRO_DIR / "scripts" / "run_d11.sh")
    copy_if_exists(SCRIPTS_DIR / "run_d12.sh", REPRO_DIR / "scripts" / "run_d12.sh")

    seed_text = "\n".join(str(s) for s in SEEDS) + "\n"
    (REPRO_DIR / "configs" / "seeds.txt").write_text(seed_text, encoding="utf-8")


def copy_frozen_splits() -> None:
    dst = REPRO_DIR / "data" / "splits"
    ensure_dir(dst)

    copy_if_exists(DATA_PROCESSED_DIR / "adult", dst / "adult")
    copy_if_exists(DATA_PROCESSED_DIR / "credit_default", dst / "credit_default")


def copy_results() -> None:
    baseline_dst = REPRO_DIR / "results" / "baseline"
    ensure_dir(baseline_dst)

    baseline_files = [
        "all_results.csv",
        "per_seed_results.csv",
        "summary.csv",
        "baseline_results.csv",
        "baseline_results_numeric.csv",
        "score_outputs.csv",
    ]

    for name in baseline_files:
        copy_if_exists(RESULTS_DIR / name, baseline_dst / name)

    copy_if_exists(RESULTS_DIR / "cards", REPRO_DIR / "results" / "cards")
    copy_if_exists(RESULTS_DIR / "figures", REPRO_DIR / "results" / "figures")
    copy_if_exists(RESULTS_DIR / "d8", REPRO_DIR / "results" / "d8_error_audit")

    copy_if_exists(D9_DIR, REPRO_DIR / "results" / "fairness" / "d9_pareto")
    copy_if_exists(D10_DIR, REPRO_DIR / "results" / "fairness" / "d10_fairness_aware")
    copy_if_exists(D11_DIR, REPRO_DIR / "results" / "ablation" / "d11_ablation")

    # LOF tuning artifact for Credit Default hyperparameter refinement.
    copy_if_exists(RESULTS_DIR / "lof_tuning_credit_default.csv", REPRO_DIR / "results" / "lof_tuning_credit_default.csv")
    copy_if_exists(RESULTS_DIR / "LOF_Tuning_Credit_Default_Report.md", REPRO_DIR / "results" / "LOF_Tuning_Credit_Default_Report.md")
    copy_if_exists(RESULTS_DIR / "LOF_Tuning_Credit_Default_Report.docx", REPRO_DIR / "results" / "LOF_Tuning_Credit_Default_Report.docx")


def copy_reports() -> None:
    reports_dst = REPRO_DIR / "reports"
    ensure_dir(reports_dst)

    report_files = [
        "D1_Problem_Statement_RQs_Hypotheses.docx",
        "D1_Problem_Statement_RQs_Hypotheses.md",
        "D2_Decision_Log.md",
        "D5_Metrics_Definition_Table.docx",
        "D5_Metrics_Definition_Table.md",
        "D8_Error_Audit_Report.docx",
        "D8_Error_Audit_Report.md",
        "D9_Pareto_Front_Report.docx",
        "D9_Pareto_Front_Report.md",
        "D10_Fairness_Aware_Results_Report.docx",
        "D10_Fairness_Aware_Results_Report.md",
        "D11_Ablation_Study_Report.docx",
        "D11_Ablation_Study_Report.md",
        "LOF_Tuning_Credit_Default_Report.docx",
        "LOF_Tuning_Credit_Default_Report.md",
    ]

    for name in report_files:
        copy_if_exists(ROOT / name, reports_dst / name)


def safe_wilcoxon_rel(a: np.ndarray, b: np.ndarray) -> float:
    """Paired Wilcoxon signed-rank p-value. Dùng cho n seed nhỏ, không giả định chuẩn."""
    try:
        from scipy import stats

        a = np.asarray(a, dtype=float)
        b = np.asarray(b, dtype=float)
        mask = np.isfinite(a) & np.isfinite(b)
        a = a[mask]
        b = b[mask]

        if len(a) < 2 or len(b) < 2 or len(a) != len(b):
            return np.nan
        if np.allclose(a, b):
            return 1.0

        return float(stats.wilcoxon(a, b, zero_method="wilcox", alternative="two-sided", mode="auto").pvalue)
    except Exception:
        return np.nan


def add_holm_correction(df: pd.DataFrame, alpha: float = 0.05) -> pd.DataFrame:
    """Thêm p-value Holm-Bonferroni cho các kiểm định non-baseline."""
    df = df.copy()
    df["p_value_holm"] = np.nan
    df["significant_holm_0_05"] = ""
    mask = df["p_value_vs_baseline"].notna()
    pvals = df.loc[mask, "p_value_vs_baseline"].astype(float)
    if pvals.empty:
        return df

    order = np.argsort(pvals.to_numpy())
    sorted_idx = pvals.index.to_numpy()[order]
    sorted_p = pvals.to_numpy()[order]
    m = len(sorted_p)
    adj = np.empty(m, dtype=float)
    running = 0.0
    for i, p in enumerate(sorted_p):
        val = (m - i) * p
        running = max(running, val)
        adj[i] = min(running, 1.0)
    df.loc[sorted_idx, "p_value_holm"] = adj
    df.loc[mask, "significant_holm_0_05"] = np.where(df.loc[mask, "p_value_holm"].astype(float) < alpha, "Yes", "No")
    return df


def make_d12_stat_summary() -> pd.DataFrame:
    """
    Tạo bảng thống kê cuối cùng cho D12:
    - baseline;
    - post-processing;
    - in-processing;
    - mean, std, change_vs_baseline, p_value_vs_baseline.
    """
    out_path = REPRO_DIR / "results" / "D12_final_statistical_summary.csv"
    ensure_dir(out_path.parent)

    all_results_path = RESULTS_DIR / "all_results.csv"
    d9_compare_path = D9_DIR / "D9_baseline_vs_recommended_postprocessing.csv"
    d9_candidates_path = D9_DIR / "d9_postprocessing_candidates.csv"
    d10_in_raw_path = D10_DIR / "D10_inprocessing_raw_results.csv"

    rows: List[Dict] = []

    if not all_results_path.exists() or not d9_compare_path.exists():
        fallback = D10_DIR / "D10_fairness_aware_results.csv"
        if fallback.exists():
            df = pd.read_csv(fallback)
            df.to_csv(out_path, index=False, encoding="utf-8-sig")
            return df

        empty = pd.DataFrame()
        empty.to_csv(out_path, index=False, encoding="utf-8-sig")
        return empty

    all_results = pd.read_csv(all_results_path)
    d9_compare = pd.read_csv(d9_compare_path)

    d9_candidates = pd.read_csv(d9_candidates_path) if d9_candidates_path.exists() else pd.DataFrame()
    d10_in_raw = pd.read_csv(d10_in_raw_path) if d10_in_raw_path.exists() else pd.DataFrame()

    metrics = ["pr_auc", "f1", "eo_gap"]

    for _, rec in d9_compare.iterrows():
        dataset = rec["dataset"]
        base_model = rec["base_model"]

        base = all_results[
            (all_results["dataset"] == dataset)
            & (all_results["model"] == base_model)
        ].copy()

        if base.empty:
            continue

        for metric in metrics:
            rows.append(
                {
                    "dataset": dataset,
                    "method": "Baseline",
                    "method_detail": str(base_model),
                    "metric": metric,
                    "mean": float(base[metric].mean()),
                    "std": float(base[metric].std()),
                    "n_seeds": int(base["seed"].nunique()) if "seed" in base.columns else len(base),
                    "change_vs_baseline": 0.0,
                    "p_value_vs_baseline": np.nan,
                }
            )

        if not d9_candidates.empty:
            rule = str(rec["recommended_rule"])
            param = float(rec["recommended_param"])

            post = d9_candidates[
                (d9_candidates["dataset"] == dataset)
                & (d9_candidates["base_model"] == base_model)
                & (d9_candidates["rule"] == rule)
            ].copy()

            if not post.empty:
                post["param_diff"] = (post["param"].astype(float) - param).abs()
                min_diff = float(post["param_diff"].min())
                post = post[post["param_diff"] <= min_diff + 1e-12].copy()

                for metric in metrics:
                    if "seed" in post.columns and "seed" in base.columns:
                        merged = base[["seed", metric]].merge(
                            post[["seed", metric]],
                            on="seed",
                            suffixes=("_base", "_post"),
                        )
                    else:
                        merged = pd.DataFrame()

                    change = (
                        float(merged[f"{metric}_post"].mean() - merged[f"{metric}_base"].mean())
                        if not merged.empty
                        else np.nan
                    )

                    p_value = (
                        safe_wilcoxon_rel(
                            merged[f"{metric}_base"].to_numpy(),
                            merged[f"{metric}_post"].to_numpy(),
                        )
                        if not merged.empty
                        else np.nan
                    )

                    rows.append(
                        {
                            "dataset": dataset,
                            "method": "Post-processing",
                            "method_detail": f"{rule} - param={param:.4f}",
                            "metric": metric,
                            "mean": float(post[metric].mean()),
                            "std": float(post[metric].std()),
                            "n_seeds": int(post["seed"].nunique()) if "seed" in post.columns else len(post),
                            "change_vs_baseline": change,
                            "p_value_vs_baseline": p_value,
                        }
                    )

        if not d10_in_raw.empty:
            in_raw = d10_in_raw[d10_in_raw["dataset"] == dataset].copy()

            if not in_raw.empty:
                for metric in metrics:
                    if "seed" in in_raw.columns and "seed" in base.columns:
                        merged = base[["seed", metric]].merge(
                            in_raw[["seed", metric]],
                            on="seed",
                            suffixes=("_base", "_in"),
                        )
                    else:
                        merged = pd.DataFrame()

                    change = (
                        float(merged[f"{metric}_in"].mean() - merged[f"{metric}_base"].mean())
                        if not merged.empty
                        else np.nan
                    )

                    p_value = (
                        safe_wilcoxon_rel(
                            merged[f"{metric}_base"].to_numpy(),
                            merged[f"{metric}_in"].to_numpy(),
                        )
                        if not merged.empty
                        else np.nan
                    )

                    rows.append(
                        {
                            "dataset": dataset,
                            "method": "In-processing",
                            "method_detail": "Reweighted IsolationForest",
                            "metric": metric,
                            "mean": float(in_raw[metric].mean()),
                            "std": float(in_raw[metric].std()),
                            "n_seeds": int(in_raw["seed"].nunique()) if "seed" in in_raw.columns else len(in_raw),
                            "change_vs_baseline": change,
                            "p_value_vs_baseline": p_value,
                        }
                    )

    df = pd.DataFrame(rows)
    if not df.empty and "p_value_vs_baseline" in df.columns:
        df = add_holm_correction(df)
        # Giữ cột kỹ thuật p_value_vs_baseline, đồng thời thêm hai cột rõ nghĩa
        # để đưa thẳng vào Bảng 4.10 trong luận văn. Hai cột này đặt cạnh nhau
        # đúng yêu cầu: p-value gốc và p-value đã hiệu chỉnh Holm-Bonferroni.
        if "p_value_raw" not in df.columns:
            insert_after = list(df.columns).index("p_value_vs_baseline") + 1
            df.insert(insert_after, "p_value_raw", df["p_value_vs_baseline"])
        if "p_value_holm_corrected" not in df.columns:
            insert_after = list(df.columns).index("p_value_raw") + 1
            df.insert(insert_after, "p_value_holm_corrected", df["p_value_holm"])
        df["correction_method"] = df["p_value_vs_baseline"].apply(
            lambda x: "Holm-Bonferroni over 12 non-baseline Wilcoxon tests" if pd.notna(x) else ""
        )
    df.to_csv(out_path, index=False, encoding="utf-8-sig")
    return df


def make_artifact_checklist() -> pd.DataFrame:
    checklist_path = REPRO_DIR / "D12_artifact_checklist.csv"
    checklist_path.touch(exist_ok=True)

    checks = [
        ("D1", ROOT / "D1_Problem_Statement_RQs_Hypotheses.docx"),
        ("D2", ROOT / "D2_Decision_Log.md"),
        ("D3_config", ROOT / "config" / "config.yaml"),
        ("D3_run_all", ROOT / "scripts" / "run_all.sh"),
        ("D4_dataset_cards", RESULTS_DIR / "cards"),
        ("D5_metrics_table", ROOT / "D5_Metrics_Definition_Table.docx"),
        ("D6_all_results", RESULTS_DIR / "all_results.csv"),
        ("D7_figures", RESULTS_DIR / "figures"),
        ("D8_error_audit", RESULTS_DIR / "d8"),
        ("D9_pareto", RESULTS_DIR / "d9"),
        ("D10_fairness_results", RESULTS_DIR / "d10"),
        ("D11_ablation", RESULTS_DIR / "d11"),
        ("D12_reproducibility", REPRO_DIR),
        ("D12_readme", REPRO_DIR / "README.md"),
        ("D12_manifest", REPRO_DIR / "MANIFEST_sha256.csv"),
        ("D12_checklist", REPRO_DIR / "D12_artifact_checklist.csv"),
        ("D12_stat_summary", REPRO_DIR / "results" / "D12_final_statistical_summary.csv"),
        ("D12_report_md", REPRO_DIR / "D12_Reproducibility_Report.md"),
    ]

    rows = []

    for name, path in checks:
        try:
            rel = path.relative_to(ROOT).as_posix()
        except Exception:
            rel = str(path)

        rows.append(
            {
                "item": name,
                "path": rel,
                "exists": bool(path.exists()),
            }
        )

    df = pd.DataFrame(rows)
    df.to_csv(checklist_path, index=False, encoding="utf-8-sig")
    return df


def write_repro_readme() -> None:
    lines: List[str] = []

    lines.append("# D12 - Reproducibility Package")
    lines.append("")
    lines.append("## 1. Mục tiêu")
    lines.append("")
    lines.append(
        "Gói này dùng để tái lập kết quả thực nghiệm cho luận văn: "
        "**Nghiên cứu bài toán phát hiện bất thường và sự công bằng trong mất cân bằng dữ liệu học máy**."
    )
    lines.append("")
    lines.append("D12 đóng gói các thành phần cần thiết để tái lập kết quả từ D1 đến D11.")
    lines.append("")
    lines.append("## 2. Cấu trúc gói D12")
    lines.append("")
    lines.append("```text")
    lines.append("reproducibility/")
    lines.append("├── configs/")
    lines.append("│   ├── config.yaml")
    lines.append("│   └── seeds.txt")
    lines.append("├── data/")
    lines.append("│   └── splits/")
    lines.append("├── results/")
    lines.append("│   ├── baseline/")
    lines.append("│   ├── fairness/")
    lines.append("│   ├── ablation/")
    lines.append("│   ├── figures/")
    lines.append("│   └── D12_final_statistical_summary.csv")
    lines.append("├── scripts/")
    lines.append("│   └── run_all.sh")
    lines.append("├── reports/")
    lines.append("├── README.md")
    lines.append("├── MANIFEST_sha256.csv")
    lines.append("└── D12_artifact_checklist.csv")
    lines.append("```")
    lines.append("")
    lines.append("## 3. Môi trường")
    lines.append("")
    lines.append("Cài thư viện bằng pip:")
    lines.append("")
    lines.append("```bash")
    lines.append("pip install -r requirements.txt")
    lines.append("```")
    lines.append("")
    lines.append("Nếu dùng conda:")
    lines.append("")
    lines.append("```bash")
    lines.append("conda env create -f environment.yml")
    lines.append("conda activate fair_ad")
    lines.append("```")
    lines.append("")
    lines.append("## 4. Seeds")
    lines.append("")
    lines.append("Các seed chính thức:")
    lines.append("")
    lines.append("```text")
    lines.append(", ".join(str(s) for s in SEEDS))
    lines.append("```")
    lines.append("")
    lines.append("## 5. Cách chạy lại toàn bộ")
    lines.append("")
    lines.append("Từ thư mục gốc project, chạy:")
    lines.append("")
    lines.append("```bash")
    lines.append("bash scripts/run_all.sh")
    lines.append("```")
    lines.append("")
    lines.append("Nếu chỉ muốn tạo lại gói D12:")
    lines.append("")
    lines.append("```bash")
    lines.append("python -m src.make_d12_repro_package")
    lines.append("```")
    lines.append("")
    lines.append("## 6. File kiểm tra")
    lines.append("")
    lines.append("- `D12_artifact_checklist.csv`: kiểm tra đủ artifact chưa.")
    lines.append("- `MANIFEST_sha256.csv`: mã băm SHA256 để kiểm tra toàn vẹn file.")
    lines.append("- `results/D12_final_statistical_summary.csv`: bảng kết quả cuối cùng.")
    lines.append("")
    lines.append("## 7. Ghi chú")
    lines.append("")
    lines.append("- Train set được kiểm tra theo nguyên tắc normal-only.")
    lines.append("- Validation set dùng để chọn threshold.")
    lines.append("- Test set chỉ dùng để đánh giá cuối cùng.")
    lines.append("- Kết quả chính được báo cáo theo nhiều seed.")
    lines.append("")

    (REPRO_DIR / "README.md").write_text("\n".join(lines), encoding="utf-8")


def format_float(x) -> str:
    if pd.isna(x):
        return ""
    return f"{float(x):.4f}"


def write_d12_report(
    checklist: pd.DataFrame,
    stats: pd.DataFrame,
    manifest: pd.DataFrame,
) -> Path:
    n_ok = int(checklist["exists"].sum())
    n_all = len(checklist)

    lines: List[str] = []

    lines.append("# D12 - Reproducibility Package Report")
    lines.append("")
    lines.append("## 1. Mục tiêu")
    lines.append(
        "D12 đóng gói code, config, seeds, dữ liệu đã chia, kết quả thực nghiệm, "
        "báo cáo và manifest để hỗ trợ tái lập kết quả."
    )
    lines.append("")
    lines.append("## 2. Kết quả kiểm tra artifact")
    lines.append(f"Số mục tồn tại: **{n_ok}/{n_all}**.")
    lines.append("")
    lines.append("| item | path | exists |")
    lines.append("| --- | --- | --- |")

    for _, row in checklist.iterrows():
        lines.append(f"| {row['item']} | {row['path']} | {row['exists']} |")

    lines.append("")
    lines.append("## 3. Bảng thống kê cuối cùng")

    if stats.empty:
        lines.append("Chưa có bảng thống kê cuối cùng.")
    else:
        lines.append("Bảng 4.10 dùng kiểm định Wilcoxon signed-rank theo từng cặp seed. Có 12 kiểm định non-baseline (2 bộ dữ liệu × 2 phương pháp × 3 metric); do đó p-value gốc được hiệu chỉnh bằng Holm-Bonferroni để kiểm soát rủi ro đa so sánh.")
        lines.append("")
        lines.append("| dataset | method | detail | metric | mean | std | change_vs_baseline | p-value gốc | p-value hiệu chỉnh Holm | significant_holm_0_05 |")
        lines.append("| --- | --- | --- | --- | ---: | ---: | ---: | ---: | ---: | :---: |")

        for _, r in stats.iterrows():
            raw_p = r.get('p_value_raw', r.get('p_value_vs_baseline', np.nan))
            holm_p = r.get('p_value_holm_corrected', r.get('p_value_holm', np.nan))
            lines.append(
                f"| {r.get('dataset', '')} | "
                f"{r.get('method', '')} | "
                f"{r.get('method_detail', '')} | "
                f"{r.get('metric', '')} | "
                f"{format_float(r.get('mean', np.nan))} | "
                f"{format_float(r.get('std', np.nan))} | "
                f"{format_float(r.get('change_vs_baseline', np.nan))} | "
                f"{format_float(raw_p)} | "
                f"{format_float(holm_p)} | "
                f"{r.get('significant_holm_0_05', '')} |"
            )

    lines.append("")
    lines.append("## 4. Manifest")
    lines.append(f"Tổng số file trong gói D12: **{len(manifest)}**.")
    lines.append("Manifest lưu tại `MANIFEST_sha256.csv`.")
    lines.append("")
    lines.append("## 5. Kết luận")
    lines.append(
        "Nếu toàn bộ các dòng quan trọng trong `D12_artifact_checklist.csv` có `exists=True`, "
        "gói D12 đạt yêu cầu kỹ thuật cơ bản cho khả năng tái lập. Các kết luận thống kê chính nên ưu tiên đọc theo cột p_value_holm để tránh diễn giải quá mức khi có nhiều phép kiểm định."
    )

    report_path = REPRO_DIR / "D12_Reproducibility_Report.md"
    report_path.write_text("\n".join(lines), encoding="utf-8")

    copy_if_exists(report_path, ROOT / "D12_Reproducibility_Report.md")

    return report_path


def try_write_docx(md_path: Path, docx_path: Path) -> None:
    try:
        from docx import Document
        from docx.shared import Pt
    except Exception:
        print("[WARN] Chưa cài python-docx nên chỉ tạo báo cáo .md.")
        return

    doc = Document()
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(13)

    for line in md_path.read_text(encoding="utf-8").splitlines():
        line = line.strip()

        if not line:
            doc.add_paragraph("")
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.startswith("|"):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = "Courier New"
            run.font.size = Pt(8)
        else:
            doc.add_paragraph(line)

    doc.save(docx_path)
    copy_if_exists(docx_path, ROOT / "D12_Reproducibility_Report.docx")


def main() -> None:
    if REPRO_DIR.exists():
        shutil.rmtree(REPRO_DIR)

    ensure_dir(REPRO_DIR)

    print("[D12] Copy configs and scripts...")
    copy_configs_and_scripts()

    print("[D12] Copy frozen splits...")
    copy_frozen_splits()

    print("[D12] Copy results...")
    copy_results()

    print("[D12] Copy reports...")
    copy_reports()

    print("[D12] Make final statistical summary...")
    stats = make_d12_stat_summary()

    print("[D12] Write README...")
    write_repro_readme()

    print("[D12] Create initial manifest...")
    manifest = make_manifest()

    print("[D12] Create artifact checklist...")
    checklist = make_artifact_checklist()

    print("[D12] Write D12 report...")
    report_md = write_d12_report(checklist, stats, manifest)
    try_write_docx(report_md, REPRO_DIR / "D12_Reproducibility_Report.docx")

    print("[D12] Refresh manifest and checklist...")
    manifest = make_manifest()
    checklist = make_artifact_checklist()
    report_md = write_d12_report(checklist, stats, manifest)
    try_write_docx(report_md, REPRO_DIR / "D12_Reproducibility_Report.docx")

    print("[D12] Final manifest...")
    manifest = make_manifest()

    print("[OK] Đã tạo xong D12.")
    print(f"[OK] Gói D12 nằm tại: {REPRO_DIR}")
    print("[OK] Các file chính:")
    print(" - reproducibility/README.md")
    print(" - reproducibility/D12_artifact_checklist.csv")
    print(" - reproducibility/MANIFEST_sha256.csv")
    print(" - reproducibility/results/D12_final_statistical_summary.csv")
    print(" - reproducibility/D12_Reproducibility_Report.md")
    print(" - reproducibility/D12_Reproducibility_Report.docx nếu đã cài python-docx")


if __name__ == "__main__":
    main()
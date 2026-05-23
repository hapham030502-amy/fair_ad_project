from __future__ import annotations

from pathlib import Path
from typing import Dict, List, Tuple

import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
from sklearn.metrics import average_precision_score, f1_score
from sklearn.ensemble import IsolationForest
from sklearn.svm import OneClassSVM

from src.run_all_models import classical_scores, run_autoencoder, set_global_seed, get_lof_n_neighbors
from src.models.deep_svdd import train_deep_svdd, score_deep_svdd
from src.models.fast_lof import FastKDTreeLOF


ROOT = Path(__file__).resolve().parents[1]

CONFIG_PATH = ROOT / "config" / "config.yaml"
PROCESSED_ROOT = ROOT / "data" / "processed"
RESULTS_DIR = ROOT / "results"
D9_DIR = RESULTS_DIR / "d9"
D11_DIR = RESULTS_DIR / "d11"

SCORE_OUTPUTS_FILE = RESULTS_DIR / "score_outputs.csv"
D9_COMPARE_FILE = D9_DIR / "D9_baseline_vs_recommended_postprocessing.csv"
D9_BEST_RULE_FILE = D9_DIR / "D9_best_tradeoff_table.csv"


GROUP_IMBALANCE_LEVELS = [0.50, 0.70, 0.80, 0.90, 0.95]
LABEL_NOISE_LEVELS = [0.00, 0.05, 0.10, 0.15]
DEFAULT_10_SEEDS = [42, 123, 456, 789, 1011, 2026, 31415, 27182, 16180, 14142]


def safe_pr_auc(y_true: np.ndarray, score: np.ndarray) -> float:
    y_true = np.asarray(y_true, dtype=int).ravel()
    score = np.asarray(score, dtype=float).ravel()

    if len(np.unique(y_true)) < 2:
        return 0.0

    return float(average_precision_score(y_true, score))


def percentile_rank(scores: np.ndarray) -> np.ndarray:
    scores = np.asarray(scores, dtype=float).ravel()

    if len(scores) <= 1:
        return np.zeros(len(scores), dtype=float)

    order = np.argsort(scores)
    ranks = np.empty(len(scores), dtype=float)
    ranks[order] = np.arange(len(scores), dtype=float)

    return ranks / (len(scores) - 1)


def compute_group_fpr_fnr(
    y_true: np.ndarray,
    y_pred: np.ndarray,
    mask: np.ndarray,
) -> Tuple[float, float]:
    yt = y_true[mask]
    yp = y_pred[mask]

    fp = int(((yp == 1) & (yt == 0)).sum())
    tn = int(((yp == 0) & (yt == 0)).sum())
    fn = int(((yp == 0) & (yt == 1)).sum())
    tp = int(((yp == 1) & (yt == 1)).sum())

    fpr = fp / (fp + tn) if (fp + tn) > 0 else 0.0
    fnr = fn / (fn + tp) if (fn + tp) > 0 else 0.0

    return float(fpr), float(fnr)


def compute_fairness(
    y_true: np.ndarray,
    y_pred: np.ndarray,
    sensitive: np.ndarray,
) -> Dict[str, float]:
    y_true = np.asarray(y_true, dtype=int).ravel()
    y_pred = np.asarray(y_pred, dtype=int).ravel()
    sensitive = np.asarray(sensitive).ravel()

    groups = np.unique(sensitive)

    if len(groups) < 2:
        return {
            "delta_fpr": 0.0,
            "delta_fnr": 0.0,
            "eo_gap": 0.0,
        }

    fprs: List[float] = []
    fnrs: List[float] = []

    for g in groups:
        fpr, fnr = compute_group_fpr_fnr(y_true, y_pred, sensitive == g)
        fprs.append(fpr)
        fnrs.append(fnr)

    delta_fpr = float(max(fprs) - min(fprs))
    delta_fnr = float(max(fnrs) - min(fnrs))

    return {
        "delta_fpr": delta_fpr,
        "delta_fnr": delta_fnr,
        "eo_gap": delta_fpr + delta_fnr,
    }


def threshold_for_target_fpr(
    y_true_g: np.ndarray,
    score_g: np.ndarray,
    target_fpr: float,
) -> float:
    y_true_g = np.asarray(y_true_g, dtype=int).ravel()
    score_g = np.asarray(score_g, dtype=float).ravel()

    normal_scores = score_g[y_true_g == 0]

    if len(normal_scores) == 0:
        normal_scores = score_g

    target_fpr = min(max(float(target_fpr), 0.0), 1.0)
    percentile = 100.0 * (1.0 - target_fpr)

    return float(np.percentile(normal_scores, percentile))


def apply_threshold_rule(
    y_true: np.ndarray,
    score: np.ndarray,
    sensitive: np.ndarray,
    rule: str,
    param: float,
) -> Tuple[np.ndarray, np.ndarray]:
    """
    Trả về:
    - y_pred: nhãn dự đoán anomaly.
    - adjusted_score: score sau hậu xử lý để tính PR-AUC.
    """
    y_true = np.asarray(y_true, dtype=int).ravel()
    score = np.asarray(score, dtype=float).ravel()
    sensitive = np.asarray(sensitive).ravel()

    rule = str(rule).strip()

    if rule == "Global threshold":
        theta = float(np.percentile(score, float(param)))
        y_pred = (score >= theta).astype(int)
        adjusted_score = score.copy()
        return y_pred, adjusted_score

    if rule == "Per-group FPR threshold":
        y_pred = np.zeros(len(score), dtype=int)
        adjusted_score = np.zeros(len(score), dtype=float)

        for g in np.unique(sensitive):
            mask = sensitive == g
            theta_g = threshold_for_target_fpr(
                y_true[mask],
                score[mask],
                target_fpr=float(param),
            )
            y_pred[mask] = (score[mask] >= theta_g).astype(int)
            adjusted_score[mask] = score[mask] - theta_g

        return y_pred, adjusted_score

    if rule == "Top-k per group":
        rank_score = np.zeros(len(score), dtype=float)

        for g in np.unique(sensitive):
            mask = sensitive == g
            rank_score[mask] = percentile_rank(score[mask])

        y_pred = (rank_score >= (1.0 - float(param))).astype(int)
        adjusted_score = rank_score
        return y_pred, adjusted_score

    raise ValueError(f"Không hỗ trợ threshold rule: {rule}")


def evaluate_setting(
    y_true: np.ndarray,
    score: np.ndarray,
    sensitive: np.ndarray,
    rule: str,
    param: float,
) -> Dict[str, float]:
    y_pred, adjusted_score = apply_threshold_rule(
        y_true=y_true,
        score=score,
        sensitive=sensitive,
        rule=rule,
        param=param,
    )

    out = {
        "pr_auc": safe_pr_auc(y_true, adjusted_score),
        "f1": float(f1_score(y_true, y_pred, zero_division=0)),
    }

    out.update(compute_fairness(y_true, y_pred, sensitive))

    return out


def resample_group_imbalance(
    df: pd.DataFrame,
    majority_ratio: float,
    seed: int,
) -> pd.DataFrame:
    """
    Tạo dữ liệu đánh giá có tỉ lệ nhóm majority theo các mức:
    50-50, 70-30, 80-20, 90-10, 95-5.

    Đây là ablation dạng controlled resampling trên test score đã có.
    """
    rng = np.random.default_rng(seed)

    groups = df["sensitive"].unique()

    if len(groups) < 2:
        return df.copy()

    counts = df["sensitive"].value_counts()
    majority_group = counts.idxmax()
    minority_group = counts.idxmin()

    maj_df = df[df["sensitive"] == majority_group]
    min_df = df[df["sensitive"] == minority_group]

    r = float(majority_ratio)
    r = min(max(r, 0.50), 0.99)

    # Ưu tiên giữ tối đa minority, tính số majority cần lấy.
    n_min = len(min_df)
    n_maj_needed = int(round((r / (1.0 - r)) * n_min))

    if n_maj_needed <= len(maj_df):
        sampled_min = min_df
        sampled_maj = maj_df.sample(
            n=n_maj_needed,
            replace=False,
            random_state=seed,
        )
    else:
        sampled_maj = maj_df
        n_min_needed = int(round(((1.0 - r) / r) * len(maj_df)))
        n_min_needed = max(1, min(n_min_needed, len(min_df)))
        sampled_min = min_df.sample(
            n=n_min_needed,
            replace=False,
            random_state=seed,
        )

    out = pd.concat([sampled_maj, sampled_min], ignore_index=True)
    out = out.sample(frac=1.0, random_state=seed).reset_index(drop=True)

    return out



def get_config_seeds_10() -> List[int]:
    """
    Lấy seed từ config/config.yaml; chuẩn hóa danh sách chính thức lên 10 seed để bảo đảm thống nhất D6-D12.

    Yêu cầu của GVHD: tăng số seed lên 10 để giảm variance cho A2.
    """
    try:
        import yaml
        if CONFIG_PATH.exists():
            with open(CONFIG_PATH, "r", encoding="utf-8-sig") as f:
                cfg = yaml.safe_load(f) or {}
            seeds = list(cfg.get("reproducibility", {}).get("seeds", []))
            seeds = [int(x) for x in seeds]
            if len(seeds) >= 10:
                return seeds[:10]
    except Exception:
        pass

    return DEFAULT_10_SEEDS


def load_processed_arrays_for_a2(
    dataset: str,
    base_model: str,
    score_outputs: pd.DataFrame,
) -> Dict[str, np.ndarray]:
    """
    Đọc dữ liệu train/val/test gốc cho A2.

    Điểm sửa quan trọng:
    - A2 KHÔNG lấy y_true từ test rồi flip nữa.
    - y_test chỉ dùng làm ground-truth sạch khi tính metric.
    - sensitive test lấy từ score_outputs.csv để giữ đúng thứ tự sample_index của kết quả cũ.
    """
    npz_path = PROCESSED_ROOT / dataset / "transformed.npz"
    if not npz_path.exists():
        raise FileNotFoundError(f"Không tìm thấy dữ liệu processed: {npz_path}")

    data = np.load(npz_path, allow_pickle=True)
    required = {"X_train", "X_val", "X_test", "y_train", "y_val", "y_test"}
    missing = required - set(data.files)
    if missing:
        raise ValueError(f"{npz_path} thiếu key: {sorted(missing)}")

    X_train = np.asarray(data["X_train"], dtype=np.float32)
    X_val = np.asarray(data["X_val"], dtype=np.float32)
    X_test = np.asarray(data["X_test"], dtype=np.float32)
    y_train = np.asarray(data["y_train"], dtype=int).ravel()
    y_val = np.asarray(data["y_val"], dtype=int).ravel()
    y_test = np.asarray(data["y_test"], dtype=int).ravel()

    if int(y_train.sum()) != 0:
        raise ValueError(
            f"A2 yêu cầu train gốc normal-only, nhưng {dataset} có {int(y_train.sum())} anomaly trong y_train."
        )

    sub = score_outputs[
        (score_outputs["dataset"] == dataset)
        & (score_outputs["model"] == base_model)
    ].copy()
    if sub.empty:
        raise ValueError(f"Không tìm thấy score_outputs cho dataset={dataset}, model={base_model}")

    first_seed = sorted(sub["seed"].unique())[0]
    sub_first = sub[sub["seed"] == first_seed].copy()
    if "sample_index" in sub_first.columns:
        sub_first = sub_first.sort_values("sample_index")

    s_test = sub_first["sensitive"].to_numpy(dtype=int)
    y_from_score_file = sub_first["y_true"].to_numpy(dtype=int)

    if len(s_test) != len(y_test):
        raise ValueError(
            f"Không khớp số mẫu test của {dataset}: s_test={len(s_test)}, y_test={len(y_test)}"
        )

    if not np.array_equal(y_from_score_file, y_test):
        raise ValueError(
            f"y_test trong transformed.npz không khớp y_true trong score_outputs.csv của {dataset}. "
            "Dừng để tránh đánh giá sai."
        )

    return {
        "X_train": X_train,
        "y_train": y_train,
        "X_val": X_val,
        "y_val": y_val,
        "X_test": X_test,
        "y_test_clean": y_test.copy(),
        "s_test": s_test,
    }


def make_train_noise_by_contamination(
    X_train_clean: np.ndarray,
    X_val: np.ndarray,
    y_val: np.ndarray,
    noise_rate: float,
    seed: int,
) -> Tuple[np.ndarray, np.ndarray, float, int]:
    """
    Mô phỏng label noise CHỈ TRÊN TRAIN cho anomaly detection.

    Vì train gốc là normal-only, nhiễu nhãn được mô phỏng bằng cách lấy một phần mẫu anomaly
    từ validation và đưa vào train như thể chúng bị gán nhãn normal. Test set tuyệt đối không bị sửa.

    Trả về:
    - X_train_noisy: train sau khi bị nhiễu.
    - cal_mask: mask validation còn lại dùng để chọn threshold.
    - actual_noise: tỷ lệ anomaly bị đưa nhầm vào train.
    - n_noisy: số anomaly thêm vào train.
    """
    noise_rate = float(noise_rate)
    rng = np.random.default_rng(seed)

    if noise_rate <= 0:
        return X_train_clean.copy(), np.ones(len(X_val), dtype=bool), 0.0, 0

    anomaly_idx = np.flatnonzero(np.asarray(y_val, dtype=int).ravel() == 1)
    if len(anomaly_idx) == 0:
        return X_train_clean.copy(), np.ones(len(X_val), dtype=bool), 0.0, 0

    # Tỷ lệ nhiễu sau khi thêm: n_noisy / (n_train + n_noisy) ~= noise_rate
    n_noisy = int(round((noise_rate / max(1e-12, 1.0 - noise_rate)) * len(X_train_clean)))
    n_noisy = max(1, n_noisy)

    replace = n_noisy > len(anomaly_idx)
    chosen = rng.choice(anomaly_idx, size=n_noisy, replace=replace)

    cal_mask = np.ones(len(X_val), dtype=bool)
    cal_mask[np.unique(chosen)] = False

    X_train_noisy = np.vstack([X_train_clean, X_val[chosen]])
    actual_noise = n_noisy / len(X_train_noisy)

    return X_train_noisy, cal_mask, float(actual_noise), int(n_noisy)


def fit_and_score_for_a2(
    model_name: str,
    X_train: np.ndarray,
    X_cal: np.ndarray,
    X_test: np.ndarray,
    seed: int,
) -> Tuple[np.ndarray, np.ndarray]:
    """
    Fit lại mô hình trên train bị nhiễu, rồi sinh score cho validation/test.
    """
    set_global_seed(seed)

    if model_name == "IsolationForest":
        model = IsolationForest(
            contamination="auto",
            random_state=seed,
            n_estimators=200,
            n_jobs=1,
        )
        model.fit(X_train)
        return classical_scores(model_name, model, X_cal), classical_scores(model_name, model, X_test)

    if model_name == "LOF":
        model = FastKDTreeLOF(n_neighbors=get_lof_n_neighbors("credit_default"))
        model.fit(X_train)
        return classical_scores(model_name, model, X_cal), classical_scores(model_name, model, X_test)

    if model_name == "OCSVM":
        model = OneClassSVM(kernel="rbf", nu=0.1, gamma="scale")
        model.fit(X_train)
        return classical_scores(model_name, model, X_cal), classical_scores(model_name, model, X_test)

    if model_name == "AutoEncoder":
        return run_autoencoder(X_train, X_cal, X_test, epochs=20, seed=seed)

    if model_name == "DeepSVDD":
        model, center = train_deep_svdd(X_train, epochs=20, seed=seed)
        return score_deep_svdd(model, center, X_cal), score_deep_svdd(model, center, X_test)

    raise ValueError(f"Không hỗ trợ model trong A2: {model_name}")


def apply_threshold_rule_from_calibration(
    y_cal: np.ndarray,
    score_cal: np.ndarray,
    score_test: np.ndarray,
    sensitive_test: np.ndarray,
    rule: str,
    param: float,
) -> Tuple[np.ndarray, np.ndarray]:
    """
    Áp dụng threshold rule cho A2.

    - Global threshold: chọn ngưỡng trên validation score.
    - Top-k per group: chọn top-k theo từng nhóm trên test score, không dùng y_test.
    - Per-group FPR threshold: cần sensitive validation; nếu không có thì dùng global calibration.
      Trong kết quả hiện tại A2 dùng Global threshold cho Adult và Top-k per group cho Credit Default.
    """
    rule = str(rule).strip()
    score_cal = np.asarray(score_cal, dtype=float).ravel()
    score_test = np.asarray(score_test, dtype=float).ravel()
    sensitive_test = np.asarray(sensitive_test).ravel()

    if rule == "Global threshold":
        theta = float(np.percentile(score_cal, float(param)))
        y_pred = (score_test >= theta).astype(int)
        return y_pred, score_test.copy()

    if rule == "Top-k per group":
        rank_score = np.zeros(len(score_test), dtype=float)
        for g in np.unique(sensitive_test):
            mask = sensitive_test == g
            rank_score[mask] = percentile_rank(score_test[mask])
        y_pred = (rank_score >= (1.0 - float(param))).astype(int)
        return y_pred, rank_score

    if rule == "Per-group FPR threshold":
        # Fallback an toàn khi không có sensitive validation: chọn threshold toàn cục theo target FPR trên validation.
        theta = threshold_for_target_fpr(y_cal, score_cal, target_fpr=float(param))
        y_pred = (score_test >= theta).astype(int)
        return y_pred, score_test - theta

    raise ValueError(f"Không hỗ trợ threshold rule trong A2: {rule}")


def evaluate_clean_test_for_a2(
    y_test_clean: np.ndarray,
    y_pred: np.ndarray,
    adjusted_score: np.ndarray,
    sensitive_test: np.ndarray,
) -> Dict[str, float]:
    """
    Tính metric trên y_test sạch. Đây là guardrail chính của A2.
    """
    out = {
        "pr_auc": safe_pr_auc(y_test_clean, adjusted_score),
        "f1": float(f1_score(y_test_clean, y_pred, zero_division=0)),
    }
    out.update(compute_fairness(y_test_clean, y_pred, sensitive_test))
    return out

def load_inputs() -> Tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    if not SCORE_OUTPUTS_FILE.exists():
        raise FileNotFoundError(
            f"Không tìm thấy {SCORE_OUTPUTS_FILE}. Cần có D6-D7 trước."
        )

    if not D9_COMPARE_FILE.exists():
        raise FileNotFoundError(
            f"Không tìm thấy {D9_COMPARE_FILE}. Cần chạy D9 trước."
        )

    if not D9_BEST_RULE_FILE.exists():
        raise FileNotFoundError(
            f"Không tìm thấy {D9_BEST_RULE_FILE}. Cần chạy D9 trước."
        )

    score_outputs = pd.read_csv(SCORE_OUTPUTS_FILE)
    d9_compare = pd.read_csv(D9_COMPARE_FILE)
    d9_best = pd.read_csv(D9_BEST_RULE_FILE)

    required_cols = {
        "dataset",
        "model",
        "seed",
        "score",
        "sensitive",
        "y_true",
    }

    missing = required_cols - set(score_outputs.columns)

    if missing:
        raise ValueError(
            f"score_outputs.csv thiếu cột: {sorted(missing)}"
        )

    return score_outputs, d9_compare, d9_best


def get_recommended_rule(
    d9_compare: pd.DataFrame,
    dataset: str,
) -> Tuple[str, float, str]:
    row = d9_compare[d9_compare["dataset"] == dataset].iloc[0]

    return (
        str(row["recommended_rule"]),
        float(row["recommended_param"]),
        str(row["base_model"]),
    )


def run_a1_group_imbalance(
    score_outputs: pd.DataFrame,
    d9_compare: pd.DataFrame,
) -> pd.DataFrame:
    rows: List[Dict] = []

    for dataset in sorted(d9_compare["dataset"].unique()):
        rule, param, base_model = get_recommended_rule(d9_compare, dataset)

        data = score_outputs[
            (score_outputs["dataset"] == dataset)
            & (score_outputs["model"] == base_model)
        ].copy()

        for seed, sub in data.groupby("seed"):
            for ratio in GROUP_IMBALANCE_LEVELS:
                ab = resample_group_imbalance(
                    sub,
                    majority_ratio=ratio,
                    seed=int(seed) + int(ratio * 1000),
                )

                metrics = evaluate_setting(
                    y_true=ab["y_true"].to_numpy(dtype=int),
                    score=ab["score"].to_numpy(dtype=float),
                    sensitive=ab["sensitive"].to_numpy(),
                    rule=rule,
                    param=param,
                )

                rows.append(
                    {
                        "ablation": "A1_group_imbalance",
                        "dataset": dataset,
                        "base_model": base_model,
                        "rule": rule,
                        "param": param,
                        "setting": f"majority_ratio={ratio:.2f}",
                        "setting_value": ratio,
                        "seed": int(seed),
                        "n_samples": len(ab),
                        "group_distribution": str(
                            ab["sensitive"].value_counts(normalize=True)
                            .round(4)
                            .to_dict()
                        ),
                        **metrics,
                    }
                )

    return pd.DataFrame(rows)



def run_a2_label_noise(
    score_outputs: pd.DataFrame,
    d9_compare: pd.DataFrame,
) -> pd.DataFrame:
    """
    A2 đã sửa theo góp ý GVHD:
    - Label noise CHỈ áp dụng trên train set.
    - KHÔNG lật y_true của test khi tính metric.
    - Fit lại mô hình cho từng mức noise và từng seed.
    - Dùng 10 seed để giảm variance.
    """
    rows: List[Dict] = []
    seeds = get_config_seeds_10()

    for dataset in sorted(d9_compare["dataset"].unique()):
        rule, param, base_model = get_recommended_rule(d9_compare, dataset)

        arrays = load_processed_arrays_for_a2(
            dataset=dataset,
            base_model=base_model,
            score_outputs=score_outputs,
        )

        X_train_clean = arrays["X_train"]
        X_val = arrays["X_val"]
        y_val = arrays["y_val"]
        X_test = arrays["X_test"]
        y_test_clean = arrays["y_test_clean"]
        s_test = arrays["s_test"]

        # Fingerprint để bảo đảm y_test không bị sửa trong quá trình chạy A2.
        y_test_fingerprint = int(
            np.dot(y_test_clean.astype(int), np.arange(1, len(y_test_clean) + 1))
        )

        for seed in seeds:
            for noise in LABEL_NOISE_LEVELS:
                X_train_noisy, cal_mask, actual_noise, n_noisy = make_train_noise_by_contamination(
                    X_train_clean=X_train_clean,
                    X_val=X_val,
                    y_val=y_val,
                    noise_rate=noise,
                    seed=int(seed) + int(noise * 10000),
                )

                X_cal = X_val[cal_mask]
                y_cal = y_val[cal_mask]

                score_cal, score_test = fit_and_score_for_a2(
                    model_name=base_model,
                    X_train=X_train_noisy,
                    X_cal=X_cal,
                    X_test=X_test,
                    seed=int(seed),
                )

                y_pred, adjusted_score = apply_threshold_rule_from_calibration(
                    y_cal=y_cal,
                    score_cal=score_cal,
                    score_test=score_test,
                    sensitive_test=s_test,
                    rule=rule,
                    param=param,
                )

                assert y_test_fingerprint == int(
                    np.dot(y_test_clean.astype(int), np.arange(1, len(y_test_clean) + 1))
                ), "LỖI A2: y_test_clean đã bị thay đổi. Không được lật nhãn test."

                metrics = evaluate_clean_test_for_a2(
                    y_test_clean=y_test_clean,
                    y_pred=y_pred,
                    adjusted_score=adjusted_score,
                    sensitive_test=s_test,
                )

                rows.append(
                    {
                        "ablation": "A2_label_noise",
                        "dataset": dataset,
                        "base_model": base_model,
                        "rule": rule,
                        "param": param,
                        "setting": f"train_label_noise={noise:.2f}",
                        "setting_value": float(noise),
                        "seed": int(seed),
                        "n_samples": len(y_test_clean),
                        "n_train_clean": len(X_train_clean),
                        "n_train_noisy": len(X_train_noisy),
                        "n_calibration": len(X_cal),
                        "actual_train_noise": actual_noise,
                        "n_noisy_train_added": n_noisy,
                        "noise_applied_to": "train_only",
                        "test_y_unchanged": True,
                        "group_distribution": str(
                            pd.Series(s_test).value_counts(normalize=True).round(4).to_dict()
                        ),
                        **metrics,
                    }
                )

    return pd.DataFrame(rows)

def run_a3_threshold_rule(
    score_outputs: pd.DataFrame,
    d9_best: pd.DataFrame,
) -> pd.DataFrame:
    rows: List[Dict] = []

    for _, best_row in d9_best.iterrows():
        dataset = str(best_row["dataset"])
        base_model = str(best_row["base_model"])
        rule = str(best_row["rule"])
        param = float(best_row["param"])

        data = score_outputs[
            (score_outputs["dataset"] == dataset)
            & (score_outputs["model"] == base_model)
        ].copy()

        for seed, sub in data.groupby("seed"):
            metrics = evaluate_setting(
                y_true=sub["y_true"].to_numpy(dtype=int),
                score=sub["score"].to_numpy(dtype=float),
                sensitive=sub["sensitive"].to_numpy(),
                rule=rule,
                param=param,
            )

            rows.append(
                {
                    "ablation": "A3_threshold_rule",
                    "dataset": dataset,
                    "base_model": base_model,
                    "rule": rule,
                    "param": param,
                    "setting": rule,
                    "setting_value": param,
                    "seed": int(seed),
                    "n_samples": len(sub),
                    "group_distribution": str(
                        sub["sensitive"].value_counts(normalize=True)
                        .round(4)
                        .to_dict()
                    ),
                    **metrics,
                }
            )

    return pd.DataFrame(rows)


def summarize(raw: pd.DataFrame) -> pd.DataFrame:
    summary = (
        raw.groupby(
            [
                "ablation",
                "dataset",
                "base_model",
                "rule",
                "param",
                "setting",
                "setting_value",
            ],
            as_index=False,
        )
        .agg(
            n_samples_mean=("n_samples", "mean"),
            n_seeds=("seed", "nunique"),
            actual_train_noise_mean=("actual_train_noise", "mean"),
            pr_auc_mean=("pr_auc", "mean"),
            pr_auc_std=("pr_auc", "std"),
            f1_mean=("f1", "mean"),
            f1_std=("f1", "std"),
            delta_fpr_mean=("delta_fpr", "mean"),
            delta_fpr_std=("delta_fpr", "std"),
            delta_fnr_mean=("delta_fnr", "mean"),
            delta_fnr_std=("delta_fnr", "std"),
            eo_gap_mean=("eo_gap", "mean"),
            eo_gap_std=("eo_gap", "std"),
        )
    )

    return summary


def make_conclusions(summary: pd.DataFrame) -> pd.DataFrame:
    rows: List[Dict[str, str]] = []

    for dataset in sorted(summary["dataset"].unique()):
        ds = summary[summary["dataset"] == dataset]

        # A1
        a1 = ds[ds["ablation"] == "A1_group_imbalance"].sort_values(
            "setting_value"
        )
        if not a1.empty:
            first = a1.iloc[0]
            last = a1.iloc[-1]
            rows.append(
                {
                    "dataset": dataset,
                    "ablation": "A1_group_imbalance",
                    "finding": (
                        f"Khi majority_ratio tăng từ {first['setting_value']:.2f} "
                        f"lên {last['setting_value']:.2f}, EO_gap thay đổi từ "
                        f"{first['eo_gap_mean']:.4f} lên {last['eo_gap_mean']:.4f}."
                    ),
                    "interpretation": (
                        "Nếu EO_gap tăng khi mất cân bằng nhóm lớn hơn, điều này cho thấy "
                        "fairness nhạy với phân bố nhóm trong dữ liệu đánh giá."
                    ),
                }
            )

        # A2
        a2 = ds[ds["ablation"] == "A2_label_noise"].sort_values(
            "setting_value"
        )
        if not a2.empty:
            first = a2.iloc[0]
            last = a2.iloc[-1]
            direction = "tăng" if float(last["pr_auc_mean"]) > float(first["pr_auc_mean"]) else "giảm"
            interpretation = (
                "A2 đã được sửa để label noise chỉ áp dụng trên train set; y_true của test được giữ nguyên khi tính metric. "
                "Nếu PR-AUC giảm khi train noise tăng, kết quả phù hợp kỳ vọng vì train bị contaminate làm giảm chất lượng học phân bố normal."
            )
            if dataset == "credit_default" and direction == "tăng":
                interpretation = (
                    "A2 đã được sửa để label noise chỉ áp dụng trên train set; y_true của test được giữ nguyên khi tính metric. "
                    "Trên Credit Default, PR-AUC vẫn tăng khi train noise tăng. Không nên diễn giải rằng nhiễu nhãn luôn có lợi; "
                    "đây có thể là artifact của dữ liệu, của quy tắc Top-k per group hoặc do một số anomaly được thêm vào train làm thay đổi ranking theo hướng thuận lợi. "
                    "Cần báo cáo trung thực và kiểm chứng thêm trên nhiều dataset/kiểu nhiễu khác."
                )
            rows.append(
                {
                    "dataset": dataset,
                    "ablation": "A2_label_noise",
                    "finding": (
                        f"Khi train_label_noise tăng từ {first['setting_value']:.2f} "
                        f"lên {last['setting_value']:.2f}, PR-AUC {direction} từ "
                        f"{first['pr_auc_mean']:.4f} đến {last['pr_auc_mean']:.4f}; "
                        f"EO_gap thay đổi từ {first['eo_gap_mean']:.4f} "
                        f"đến {last['eo_gap_mean']:.4f}."
                    ),
                    "interpretation": interpretation,
                }
            )

        # A3
        a3 = ds[ds["ablation"] == "A3_threshold_rule"].copy()
        if not a3.empty:
            best_eo = a3.loc[a3["eo_gap_mean"].idxmin()]
            best_pr = a3.loc[a3["pr_auc_mean"].idxmax()]
            rows.append(
                {
                    "dataset": dataset,
                    "ablation": "A3_threshold_rule",
                    "finding": (
                        f"Rule có EO_gap thấp nhất là {best_eo['rule']} "
                        f"(EO_gap={best_eo['eo_gap_mean']:.4f}); "
                        f"rule có PR-AUC cao nhất là {best_pr['rule']} "
                        f"(PR-AUC={best_pr['pr_auc_mean']:.4f})."
                    ),
                    "interpretation": (
                        "Kết quả này cho thấy lựa chọn threshold rule ảnh hưởng trực tiếp "
                        "đến trade-off utility–fairness."
                    ),
                }
            )

    return pd.DataFrame(rows)


def plot_a1(summary: pd.DataFrame, out_path: Path) -> None:
    sub = summary[summary["ablation"] == "A1_group_imbalance"]

    plt.figure(figsize=(8.5, 5.8))

    for dataset, ds in sub.groupby("dataset"):
        ds = ds.sort_values("setting_value")
        plt.plot(
            ds["setting_value"],
            ds["eo_gap_mean"],
            marker="o",
            label=dataset,
        )

    plt.xlabel("Majority group ratio")
    plt.ylabel("EO_gap")
    plt.title("D11-A1: Group Imbalance vs EO_gap")
    plt.grid(alpha=0.25)
    plt.legend()
    plt.tight_layout()
    plt.savefig(out_path, dpi=300, bbox_inches="tight")
    plt.close()


def plot_a2(summary: pd.DataFrame, out_path: Path) -> None:
    sub = summary[summary["ablation"] == "A2_label_noise"]

    plt.figure(figsize=(8.5, 5.8))

    for dataset, ds in sub.groupby("dataset"):
        ds = ds.sort_values("setting_value")
        plt.plot(
            ds["setting_value"],
            ds["pr_auc_mean"],
            marker="o",
            label=dataset,
        )

    plt.xlabel("Train label noise rate")
    plt.ylabel("PR-AUC")
    plt.title("D11-A2: Train Label Noise vs PR-AUC on Clean Test")
    plt.grid(alpha=0.25)
    plt.legend()
    plt.tight_layout()
    plt.savefig(out_path, dpi=300, bbox_inches="tight")
    plt.close()


def plot_a3(summary: pd.DataFrame, out_path: Path) -> None:
    sub = summary[summary["ablation"] == "A3_threshold_rule"].copy()

    sub["dataset_rule"] = sub["dataset"] + " | " + sub["rule"]

    plt.figure(figsize=(11.0, 6.0))
    plt.bar(
        sub["dataset_rule"],
        sub["eo_gap_mean"],
    )
    plt.xticks(rotation=35, ha="right")
    plt.ylabel("EO_gap")
    plt.title("D11-A3: Threshold Rule Comparison")
    plt.grid(axis="y", alpha=0.25)
    plt.tight_layout()
    plt.savefig(out_path, dpi=300, bbox_inches="tight")
    plt.close()


def format_float(x) -> str:
    if pd.isna(x):
        return ""
    return f"{float(x):.4f}"


def markdown_table(df: pd.DataFrame, cols: List[str]) -> str:
    d = df[cols].copy()

    for c in d.columns:
        if pd.api.types.is_float_dtype(d[c]):
            d[c] = d[c].map(format_float)

    header = "| " + " | ".join(d.columns) + " |"
    sep = "| " + " | ".join(["---"] * len(d.columns)) + " |"

    rows = []
    for _, row in d.iterrows():
        rows.append(
            "| " + " | ".join(str(row[c]) for c in d.columns) + " |"
        )

    return "\n".join([header, sep] + rows)


def write_report_md(
    summary: pd.DataFrame,
    conclusions: pd.DataFrame,
    out_path: Path,
) -> None:
    lines: List[str] = []

    lines.append("# D11 - Ablation Study")
    lines.append("")
    lines.append("## 1. Mục tiêu")
    lines.append(
        "D11 đánh giá độ nhạy của kết quả fairness-aware anomaly detection "
        "thông qua ba loại ablation: group imbalance, label noise và threshold rule."
    )
    lines.append("")
    lines.append("## 2. Thiết kế ablation")
    lines.append("- **A1 - Group Imbalance**: thay đổi tỉ lệ nhóm majority/minority theo các mức 50-50, 70-30, 80-20, 90-10, 95-5.")
    lines.append("- **A2 - Label Noise**: đã sửa theo góp ý GVHD; nhiễu nhãn chỉ áp dụng trên train set bằng cách đưa anomaly từ validation vào train như mẫu normal; y_true của test được giữ nguyên khi tính PR-AUC/F1/EO_gap; số seed dùng cho A2 là 10.")
    lines.append("- **A3 - Threshold Rule**: so sánh ba quy tắc hậu xử lý từ D9: Global threshold, Per-group FPR threshold và Top-k per group.")
    lines.append("")
    lines.append("## 3. Bảng kết quả tổng hợp")
    view = summary[
        [
            "ablation",
            "dataset",
            "rule",
            "setting",
            "pr_auc_mean",
            "f1_mean",
            "eo_gap_mean",
            "delta_fpr_mean",
            "delta_fnr_mean",
        ]
    ].copy()
    lines.append(
        markdown_table(
            view,
            [
                "ablation",
                "dataset",
                "rule",
                "setting",
                "pr_auc_mean",
                "f1_mean",
                "eo_gap_mean",
                "delta_fpr_mean",
                "delta_fnr_mean",
            ],
        )
    )
    lines.append("")
    lines.append("## 4. Kết luận theo từng ablation")
    lines.append(
        markdown_table(
            conclusions,
            [
                "dataset",
                "ablation",
                "finding",
                "interpretation",
            ],
        )
    )
    lines.append("")
    lines.append("## 5. Nhận xét")
    lines.append(
        "Ablation A1 cho biết mức độ nhạy của EO_gap khi phân bố nhóm thay đổi. "
        "Ablation A2 cho biết độ ổn định của PR-AUC và EO_gap khi train set bị nhiễu nhãn; y_true của test luôn giữ sạch. "
        "Ablation A3 cho thấy threshold rule là yếu tố quan trọng ảnh hưởng trực tiếp "
        "đến trade-off giữa utility và fairness."
    )
    lines.append("")
    lines.append("## 6. Threats to validity and limitations")
    lines.append(
        "- **Giới hạn về dữ liệu**: Thực nghiệm mới được kiểm tra trên hai bộ dữ liệu dạng bảng "
        "là Adult và Credit Default, vì vậy khả năng khái quát sang dữ liệu ảnh, chuỗi thời gian "
        "hoặc đồ thị cần được kiểm chứng thêm."
    )
    lines.append(
        "- **Giới hạn về thuộc tính nhạy cảm**: Các phân tích fairness chủ yếu sử dụng một thuộc tính "
        "nhạy cảm mặc định cho mỗi bộ dữ liệu. Các thuộc tính khác như race hoặc age có thể tạo ra "
        "kết quả khác và cần được phân tích bổ sung trong các nghiên cứu tiếp theo."
    )
    lines.append(
        "- **Giới hạn về mô hình**: Ablation sử dụng điểm bất thường đã sinh từ các mô hình baseline "
        "và các quy tắc hậu xử lý. Do đó, kết quả phản ánh độ nhạy của pipeline hiện tại, chưa đại diện "
        "cho toàn bộ các thuật toán phát hiện bất thường khác."
    )
    lines.append(
        "- **Giới hạn về nhiễu nhãn**: Thí nghiệm label noise được mô phỏng bằng cách lật nhãn ngẫu nhiên. "
        "Trong dữ liệu thực tế, nhiễu nhãn có thể có cấu trúc phức tạp hơn và phụ thuộc vào nhóm người dùng."
    )
    lines.append(
        "- **Giới hạn về threshold rule**: Các quy tắc threshold được khảo sát gồm Global threshold, "
        "Per-group FPR threshold và Top-k per group. Những chiến lược hậu xử lý phức tạp hơn có thể "
        "đem lại trade-off khác giữa utility và fairness."
    )
    lines.append("")
    lines.append("## 7. Kết luận D11")
    lines.append(
        "D11 đã tạo được bảng ablation cho ba yếu tố chính: group imbalance, label noise "
        "và threshold rule. Kết quả này dùng để viết phần robustness/ablation trong Chương 4 "
        "và làm cơ sở thảo luận ở Chương 5."
    )
    lines.append("")
    lines.append("## 8. File kết quả")
    lines.append("- `results/d11/D11_ablation_raw_results.csv`")
    lines.append("- `results/d11/D11_ablation_summary.csv`")
    lines.append("- `results/d11/D11_ablation_conclusions.csv`")
    lines.append("- `results/d11/d11_a1_group_imbalance_eogap.png`")
    lines.append("- `results/d11/d11_a2_label_noise_prauc.png`")
    lines.append("- `results/d11/d11_a3_threshold_rule_eogap.png`")
    lines.append("- `results/d11/D11_Ablation_Study_Report.md`")
    lines.append("- `results/d11/D11_Ablation_Study_Report.docx` nếu có `python-docx`")

    out_path.write_text("\n".join(lines), encoding="utf-8")


def try_write_docx(md_path: Path, docx_path: Path) -> None:
    try:
        from docx import Document
        from docx.shared import Pt
    except Exception:
        print(
            "[WARN] Chưa cài python-docx nên chỉ tạo .md. "
            "Nếu muốn có .docx, chạy: pip install python-docx"
        )
        return

    doc = Document()
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(13)

    for line in md_path.read_text(encoding="utf-8").splitlines():
        line = line.strip()

        if not line:
            doc.add_paragraph("")
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.startswith("|"):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = "Courier New"
            run.font.size = Pt(8)
        else:
            doc.add_paragraph(line)

    doc.save(docx_path)


def main() -> None:
    D11_DIR.mkdir(parents=True, exist_ok=True)

    print("[D11] Loading D9 and score outputs...")
    score_outputs, d9_compare, d9_best = load_inputs()

    print("[D11] Running A1 - Group imbalance ablation...")
    a1 = run_a1_group_imbalance(score_outputs, d9_compare)

    print("[D11] Running A2 - Label noise ablation...")
    a2 = run_a2_label_noise(score_outputs, d9_compare)

    print("[D11] Running A3 - Threshold rule ablation...")
    a3 = run_a3_threshold_rule(score_outputs, d9_best)

    raw = pd.concat([a1, a2, a3], ignore_index=True)
    raw.to_csv(D11_DIR / "D11_ablation_raw_results.csv", index=False)

    summary = summarize(raw)
    summary.to_csv(D11_DIR / "D11_ablation_summary.csv", index=False)

    conclusions = make_conclusions(summary)
    conclusions.to_csv(D11_DIR / "D11_ablation_conclusions.csv", index=False)

    print("[D11] Creating figures...")
    plot_a1(summary, D11_DIR / "d11_a1_group_imbalance_eogap.png")
    plot_a2(summary, D11_DIR / "d11_a2_label_noise_prauc.png")
    plot_a3(summary, D11_DIR / "d11_a3_threshold_rule_eogap.png")

    report_md = D11_DIR / "D11_Ablation_Study_Report.md"
    write_report_md(summary, conclusions, report_md)

    root_report_md = ROOT / "D11_Ablation_Study_Report.md"
    root_report_md.write_text(
        report_md.read_text(encoding="utf-8"),
        encoding="utf-8",
    )

    try_write_docx(
        report_md,
        D11_DIR / "D11_Ablation_Study_Report.docx",
    )
    try_write_docx(
        root_report_md,
        ROOT / "D11_Ablation_Study_Report.docx",
    )

    print("[OK] Đã tạo xong D11.")
    print(f"[OK] Kết quả nằm trong: {D11_DIR}")
    print("[OK] File chính:")
    print(" - results/d11/D11_ablation_raw_results.csv")
    print(" - results/d11/D11_ablation_summary.csv")
    print(" - results/d11/D11_ablation_conclusions.csv")
    print(" - results/d11/d11_a1_group_imbalance_eogap.png")
    print(" - results/d11/d11_a2_label_noise_prauc.png")
    print(" - results/d11/d11_a3_threshold_rule_eogap.png")
    print(" - results/d11/D11_Ablation_Study_Report.md")
    print(" - results/d11/D11_Ablation_Study_Report.docx nếu đã cài python-docx")


if __name__ == "__main__":
    main()
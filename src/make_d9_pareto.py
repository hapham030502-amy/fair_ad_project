from __future__ import annotations

from pathlib import Path
from typing import Dict, List, Tuple

import numpy as np
import pandas as pd
import matplotlib.pyplot as plt


ROOT = Path(__file__).resolve().parents[1]
RESULTS_DIR = ROOT / "results"
D9_DIR = RESULTS_DIR / "d9"

ALL_RESULTS_FILE = RESULTS_DIR / "all_results.csv"
SCORE_OUTPUTS_FILE = RESULTS_DIR / "score_outputs.csv"


def average_precision_fast(y_true: np.ndarray, scores: np.ndarray) -> float:
    """
    Tính PR-AUC theo dạng Average Precision.
    Không dùng sklearn để giảm phụ thuộc khi chạy trên máy khác.
    """
    y_true = np.asarray(y_true, dtype=int).ravel()
    scores = np.asarray(scores, dtype=float).ravel()

    total_pos = int((y_true == 1).sum())
    if total_pos == 0:
        return 0.0

    order = np.argsort(scores)[::-1]
    y_sorted = y_true[order]

    tp_cum = np.cumsum(y_sorted == 1)
    rank = np.arange(1, len(y_sorted) + 1)
    precision = tp_cum / rank

    return float(precision[y_sorted == 1].sum() / total_pos)


def f1_fast(y_true: np.ndarray, y_pred: np.ndarray) -> float:
    y_true = np.asarray(y_true, dtype=int).ravel()
    y_pred = np.asarray(y_pred, dtype=int).ravel()

    tp = int(((y_pred == 1) & (y_true == 1)).sum())
    fp = int(((y_pred == 1) & (y_true == 0)).sum())
    fn = int(((y_pred == 0) & (y_true == 1)).sum())

    denom = 2 * tp + fp + fn
    if denom == 0:
        return 0.0

    return float((2 * tp) / denom)


def fpr_fnr_fast(
    y_true: np.ndarray,
    y_pred: np.ndarray,
    mask: np.ndarray,
) -> Tuple[float, float]:
    yt = y_true[mask]
    yp = y_pred[mask]

    fp = int(((yp == 1) & (yt == 0)).sum())
    tn = int(((yp == 0) & (yt == 0)).sum())
    fn = int(((yp == 0) & (yt == 1)).sum())
    tp = int(((yp == 1) & (yt == 1)).sum())

    fpr = fp / (fp + tn) if (fp + tn) > 0 else 0.0
    fnr = fn / (fn + tp) if (fn + tp) > 0 else 0.0

    return float(fpr), float(fnr)


def fairness_fast(
    y_true: np.ndarray,
    y_pred: np.ndarray,
    sensitive: np.ndarray,
) -> Dict[str, float]:
    """
    Tính:
    - delta_fpr = max(FPR_g) - min(FPR_g)
    - delta_fnr = max(FNR_g) - min(FNR_g)
    - eo_gap = delta_fpr + delta_fnr
    """
    y_true = np.asarray(y_true, dtype=int).ravel()
    y_pred = np.asarray(y_pred, dtype=int).ravel()
    sensitive = np.asarray(sensitive).ravel()

    groups = np.unique(sensitive)

    if len(groups) < 2:
        return {
            "delta_fpr": 0.0,
            "delta_fnr": 0.0,
            "eo_gap": 0.0,
        }

    fprs: List[float] = []
    fnrs: List[float] = []

    for g in groups:
        fpr, fnr = fpr_fnr_fast(y_true, y_pred, sensitive == g)
        fprs.append(fpr)
        fnrs.append(fnr)

    delta_fpr = float(max(fprs) - min(fprs))
    delta_fnr = float(max(fnrs) - min(fnrs))

    return {
        "delta_fpr": delta_fpr,
        "delta_fnr": delta_fnr,
        "eo_gap": delta_fpr + delta_fnr,
    }


def evaluate_postprocessing(
    y_true: np.ndarray,
    y_pred: np.ndarray,
    adjusted_score: np.ndarray,
    sensitive: np.ndarray,
) -> Dict[str, float]:
    out = {
        "pr_auc": average_precision_fast(y_true, adjusted_score),
        "f1": f1_fast(y_true, y_pred),
    }
    out.update(fairness_fast(y_true, y_pred, sensitive))
    return out


def choose_base_models(all_results: pd.DataFrame) -> pd.DataFrame:
    """
    Chọn base model tốt nhất cho mỗi dataset theo mean PR-AUC từ D6.
    """
    required_cols = {"dataset", "model", "pr_auc", "eo_gap", "f1"}

    missing = required_cols - set(all_results.columns)
    if missing:
        raise ValueError(
            f"File all_results.csv thiếu cột: {sorted(missing)}"
        )

    grouped = (
        all_results
        .groupby(["dataset", "model"], as_index=False)
        .agg(
            mean_pr_auc=("pr_auc", "mean"),
            mean_eo_gap=("eo_gap", "mean"),
            mean_f1=("f1", "mean"),
        )
        .sort_values(
            ["dataset", "mean_pr_auc"],
            ascending=[True, False],
        )
    )

    idx = grouped.groupby("dataset")["mean_pr_auc"].idxmax()

    selected = grouped.loc[idx].copy().reset_index(drop=True)
    selected = selected.rename(columns={"model": "base_model"})

    return selected[
        [
            "dataset",
            "base_model",
            "mean_pr_auc",
            "mean_eo_gap",
            "mean_f1",
        ]
    ]


def threshold_for_target_fpr(
    y_true_g: np.ndarray,
    scores_g: np.ndarray,
    target_fpr: float,
) -> float:
    """
    Tìm ngưỡng theo target FPR của từng nhóm.
    Dùng phân vị của score trên các mẫu normal.
    """
    y_true_g = np.asarray(y_true_g, dtype=int)
    scores_g = np.asarray(scores_g, dtype=float)

    normal_scores = scores_g[y_true_g == 0]

    if len(normal_scores) == 0:
        normal_scores = scores_g

    target_fpr = min(max(float(target_fpr), 0.0), 1.0)
    percentile = 100.0 * (1.0 - target_fpr)

    return float(np.percentile(normal_scores, percentile))


def percentile_rank(scores: np.ndarray) -> np.ndarray:
    """
    Đổi score thành percentile rank trong từng nhóm.
    Dùng cho top-k per group.
    """
    scores = np.asarray(scores, dtype=float)

    order = np.argsort(scores)
    ranks = np.empty(len(scores), dtype=float)
    ranks[order] = np.arange(len(scores), dtype=float)

    if len(scores) <= 1:
        return np.zeros(len(scores), dtype=float)

    return ranks / (len(scores) - 1)


def is_pareto_optimal(df: pd.DataFrame) -> pd.Series:
    """
    Pareto optimal:
    - PR-AUC càng cao càng tốt.
    - EO-gap càng thấp càng tốt.
    """
    pareto = np.ones(len(df), dtype=bool)

    vals = df[["pr_auc_mean", "eo_gap_mean"]].to_numpy(dtype=float)

    for i, (pr_i, eo_i) in enumerate(vals):
        for j, (pr_j, eo_j) in enumerate(vals):
            if i == j:
                continue

            dominated = (
                pr_j >= pr_i
                and eo_j <= eo_i
                and (pr_j > pr_i or eo_j < eo_i)
            )

            if dominated:
                pareto[i] = False
                break

    return pd.Series(pareto, index=df.index)


def add_tradeoff_scores(
    mean_df: pd.DataFrame,
    base_models: pd.DataFrame,
) -> pd.DataFrame:
    base = base_models.rename(
        columns={
            "mean_pr_auc": "baseline_pr_auc_mean",
            "mean_eo_gap": "baseline_eo_gap_mean",
            "mean_f1": "baseline_f1_mean",
        }
    )

    df = mean_df.merge(
        base,
        on=["dataset", "base_model"],
        how="left",
    )

    df["pr_auc_loss_vs_baseline"] = (
        df["baseline_pr_auc_mean"] - df["pr_auc_mean"]
    )
    df["eo_gap_reduction_vs_baseline"] = (
        df["baseline_eo_gap_mean"] - df["eo_gap_mean"]
    )
    df["f1_ratio_vs_baseline"] = (
        df["f1_mean"] / df["baseline_f1_mean"].replace(0, np.nan)
    )

    df["distance_to_ideal"] = np.nan

    for dataset, idx in df.groupby("dataset").groups.items():
        sub = df.loc[list(idx)]

        pr_range = max(
            sub["pr_auc_mean"].max() - sub["pr_auc_mean"].min(),
            1e-12,
        )
        eo_range = max(
            sub["eo_gap_mean"].max() - sub["eo_gap_mean"].min(),
            1e-12,
        )

        max_pr = sub["pr_auc_mean"].max()
        min_eo = sub["eo_gap_mean"].min()

        distance = np.sqrt(
            ((max_pr - sub["pr_auc_mean"]) / pr_range) ** 2
            + ((sub["eo_gap_mean"] - min_eo) / eo_range) ** 2
        )

        df.loc[sub.index, "distance_to_ideal"] = distance

    # Guardrail: tránh chọn nghiệm F1 quá thấp
    df["passes_f1_guardrail"] = (
        df["f1_mean"] >= 0.80 * df["baseline_f1_mean"]
    )

    return df


def plot_dataset(
    df: pd.DataFrame,
    dataset: str,
    out_png: Path,
) -> None:
    sub = df[df["dataset"] == dataset].copy()

    plt.figure(figsize=(8.5, 6.2))

    rules = [
        "Global threshold",
        "Per-group FPR threshold",
        "Top-k per group",
    ]

    for rule in rules:
        r = sub[sub["rule"] == rule].sort_values("eo_gap_mean")

        if not r.empty:
            plt.plot(
                r["eo_gap_mean"],
                r["pr_auc_mean"],
                marker="o",
                markersize=3,
                linewidth=1.4,
                label=rule,
            )

    pareto = sub[sub["is_pareto"]]

    if not pareto.empty:
        plt.scatter(
            pareto["eo_gap_mean"],
            pareto["pr_auc_mean"],
            marker="*",
            s=130,
            label="Pareto-optimal",
        )

    selected = sub[sub["is_selected_per_rule"]]

    if not selected.empty:
        plt.scatter(
            selected["eo_gap_mean"],
            selected["pr_auc_mean"],
            marker="D",
            s=70,
            label="Best per rule",
        )

    recommended = sub[sub["is_overall_recommended"]]

    if not recommended.empty:
        plt.scatter(
            recommended["eo_gap_mean"],
            recommended["pr_auc_mean"],
            marker="X",
            s=110,
            label="Recommended",
        )

    first = sub.iloc[0]

    plt.scatter(
        [first["baseline_eo_gap_mean"]],
        [first["baseline_pr_auc_mean"]],
        marker="x",
        s=100,
        label="Baseline",
    )

    plt.xlabel("EO_gap - thấp hơn là tốt hơn")
    plt.ylabel("PR-AUC - cao hơn là tốt hơn")
    plt.title(f"D9 Pareto Front: Utility vs Fairness - {dataset}")
    plt.grid(True, alpha=0.25)
    plt.legend(fontsize=8)
    plt.tight_layout()
    plt.savefig(out_png, dpi=300, bbox_inches="tight")
    plt.close()


def markdown_table(df: pd.DataFrame, cols: List[str]) -> str:
    d = df[cols].copy()

    for c in d.columns:
        if pd.api.types.is_float_dtype(d[c]):
            d[c] = d[c].map(
                lambda x: f"{x:.4f}" if pd.notna(x) else ""
            )

    header = "| " + " | ".join(d.columns) + " |"
    sep = "| " + " | ".join(["---"] * len(d.columns)) + " |"

    rows = []
    for _, row in d.iterrows():
        rows.append(
            "| " + " | ".join(str(row[c]) for c in d.columns) + " |"
        )

    return "\n".join([header, sep] + rows)


def write_report_md(
    best_per_rule: pd.DataFrame,
    compare: pd.DataFrame,
    out_path: Path,
) -> None:
    lines: List[str] = []

    lines.append("# D9 - Pareto Front Analysis: Utility vs Fairness cho post-processing")
    lines.append("")
    lines.append("## 1. Mục tiêu")
    lines.append(
        "D9 đánh giá trade-off giữa hiệu quả phát hiện bất thường và công bằng nhóm "
        "bằng các quy tắc hậu xử lý ngưỡng. Trục x là EO_gap, trục y là PR-AUC."
    )
    lines.append("")
    lines.append("## 2. Dữ liệu đầu vào")
    lines.append("- Kết quả baseline từ `results/all_results.csv`.")
    lines.append("- Điểm bất thường từng mẫu từ `results/score_outputs.csv`.")
    lines.append("- Mỗi dataset chọn base model có mean PR-AUC cao nhất từ D6.")
    lines.append("")
    lines.append("## 3. Các quy tắc hậu xử lý")
    lines.append("- **Global threshold**: dùng một ngưỡng chung cho toàn bộ nhóm.")
    lines.append("- **Per-group FPR threshold**: chọn ngưỡng riêng theo từng nhóm để điều chỉnh FPR.")
    lines.append("- **Top-k per group**: chọn top-k% điểm bất thường cao nhất trong từng nhóm.")
    lines.append("")
    lines.append("## 4. Best trade-off theo từng quy tắc")
    lines.append(
        markdown_table(
            best_per_rule,
            [
                "dataset",
                "base_model",
                "rule",
                "param",
                "pr_auc_mean",
                "f1_mean",
                "eo_gap_mean",
                "eo_gap_reduction_vs_baseline",
                "passes_f1_guardrail",
            ],
        )
    )
    lines.append("")
    lines.append("## 5. So sánh baseline và post-processing khuyến nghị")
    lines.append(
        markdown_table(
            compare,
            [
                "dataset",
                "base_model",
                "baseline_pr_auc",
                "baseline_f1",
                "baseline_eo_gap",
                "recommended_rule",
                "recommended_param",
                "post_pr_auc",
                "post_f1",
                "post_eo_gap",
                "eo_gap_reduction",
                "pr_auc_change",
            ],
        )
    )
    lines.append("")
    lines.append("## 6. Nhận xét")
    for _, row in compare.iterrows():
        lines.append(
            f"- **{row['dataset']}**: mô hình nền là **{row['base_model']}**. "
            f"Quy tắc khuyến nghị là **{row['recommended_rule']}** với tham số "
            f"{row['recommended_param']:.4f}; EO_gap giảm "
            f"{row['eo_gap_reduction']:.4f}, PR-AUC thay đổi "
            f"{row['pr_auc_change']:.4f}."
        )

    lines.append("")
    lines.append("## 7. Kết luận D9")
    lines.append(
        "D9 đã tạo được bảng và hình Pareto Front để phân tích trade-off giữa "
        "utility và fairness. Kết quả này là đầu vào cho D10, nơi so sánh "
        "Baseline, Post-processing và In-processing."
    )
    lines.append("")
    lines.append("## 8. File kết quả")
    lines.append("- `results/d9/d9_postprocessing_candidates.csv`")
    lines.append("- `results/d9/d9_pareto_mean.csv`")
    lines.append("- `results/d9/D9_best_tradeoff_table.csv`")
    lines.append("- `results/d9/D9_baseline_vs_recommended_postprocessing.csv`")
    lines.append("- `results/d9/pareto_adult.png`")
    lines.append("- `results/d9/pareto_credit_default.png`")

    out_path.write_text("\n".join(lines), encoding="utf-8")


def try_write_docx(md_path: Path, docx_path: Path) -> None:
    """
    Nếu máy có python-docx thì xuất thêm Word.
    Nếu chưa có python-docx, chương trình vẫn chạy và giữ file .md.
    """
    try:
        from docx import Document
        from docx.shared import Pt
    except Exception:
        print(
            "[WARN] Chưa cài python-docx nên chỉ tạo báo cáo .md. "
            "Nếu muốn xuất .docx, chạy: pip install python-docx"
        )
        return

    doc = Document()
    styles = doc.styles

    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(13)

    content = md_path.read_text(encoding="utf-8").splitlines()

    for line in content:
        line = line.strip()

        if not line:
            doc.add_paragraph("")
            continue

        if line.startswith("# "):
            doc.add_heading(line.replace("# ", ""), level=1)
        elif line.startswith("## "):
            doc.add_heading(line.replace("## ", ""), level=2)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.startswith("|"):
            # Bảng markdown giữ nguyên dạng text để tránh lỗi định dạng
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        else:
            doc.add_paragraph(line)

    doc.save(docx_path)


def check_input_files() -> None:
    if not ALL_RESULTS_FILE.exists():
        raise FileNotFoundError(
            f"Không tìm thấy {ALL_RESULTS_FILE}. "
            "Cần chạy D6 trước khi làm D9."
        )

    if not SCORE_OUTPUTS_FILE.exists():
        raise FileNotFoundError(
            f"Không tìm thấy {SCORE_OUTPUTS_FILE}. "
            "Cần chạy D6-D7 trước khi làm D9."
        )


def main() -> None:
    check_input_files()

    D9_DIR.mkdir(parents=True, exist_ok=True)

    all_results = pd.read_csv(ALL_RESULTS_FILE)
    score_outputs = pd.read_csv(SCORE_OUTPUTS_FILE)

    required_score_cols = {
        "dataset",
        "model",
        "seed",
        "score",
        "sensitive",
        "y_true",
    }

    missing_score_cols = required_score_cols - set(score_outputs.columns)

    if missing_score_cols:
        raise ValueError(
            f"File score_outputs.csv thiếu cột: {sorted(missing_score_cols)}"
        )

    base_models = choose_base_models(all_results)
    base_models.to_csv(
        D9_DIR / "D9_selected_base_models.csv",
        index=False,
    )

    rows: List[Dict[str, float]] = []

    # Lưới tham số cho 3 threshold rules
    global_grid = np.arange(70.0, 99.6, 0.5)
    target_fpr_grid = np.linspace(0.01, 0.30, 40)
    topk_grid = np.linspace(0.02, 0.30, 40)

    for _, base_row in base_models.iterrows():
        dataset = base_row["dataset"]
        model = base_row["base_model"]

        df = score_outputs[
            (score_outputs["dataset"] == dataset)
            & (score_outputs["model"] == model)
        ].copy()

        if df.empty:
            raise ValueError(
                f"Không có score_outputs cho dataset={dataset}, model={model}"
            )

        for seed, sub in df.groupby("seed"):
            y_true = sub["y_true"].to_numpy(dtype=int)
            scores = sub["score"].to_numpy(dtype=float)
            sensitive = sub["sensitive"].to_numpy()

            raw_pr_auc = average_precision_fast(y_true, scores)
            groups = np.unique(sensitive)

            # Rule 1: Global threshold
            for p in global_grid:
                theta = float(np.percentile(scores, p))
                y_pred = (scores >= theta).astype(int)

                metrics = {
                    "pr_auc": raw_pr_auc,
                    "f1": f1_fast(y_true, y_pred),
                }
                metrics.update(fairness_fast(y_true, y_pred, sensitive))

                rows.append(
                    {
                        "dataset": dataset,
                        "base_model": model,
                        "seed": int(seed),
                        "rule": "Global threshold",
                        "param": float(p),
                        "param_type": "score percentile",
                        "global_threshold": theta,
                        **metrics,
                    }
                )

            # Rule 2: Per-group FPR threshold
            for target_fpr in target_fpr_grid:
                adjusted_score = np.zeros(len(scores), dtype=float)
                y_pred = np.zeros(len(scores), dtype=int)
                extra: Dict[str, float] = {}

                for g in groups:
                    mask = sensitive == g
                    theta_g = threshold_for_target_fpr(
                        y_true[mask],
                        scores[mask],
                        target_fpr,
                    )

                    adjusted_score[mask] = scores[mask] - theta_g
                    y_pred[mask] = (scores[mask] >= theta_g).astype(int)
                    extra[f"threshold_group_{g}"] = theta_g

                metrics = evaluate_postprocessing(
                    y_true,
                    y_pred,
                    adjusted_score,
                    sensitive,
                )

                rows.append(
                    {
                        "dataset": dataset,
                        "base_model": model,
                        "seed": int(seed),
                        "rule": "Per-group FPR threshold",
                        "param": float(target_fpr),
                        "param_type": "target FPR per group",
                        **extra,
                        **metrics,
                    }
                )

            # Rule 3: Top-k per group
            rank_score = np.zeros(len(scores), dtype=float)

            for g in groups:
                mask = sensitive == g
                rank_score[mask] = percentile_rank(scores[mask])

            rank_pr_auc = average_precision_fast(y_true, rank_score)

            for k in topk_grid:
                y_pred = (rank_score >= (1.0 - float(k))).astype(int)

                metrics = {
                    "pr_auc": rank_pr_auc,
                    "f1": f1_fast(y_true, y_pred),
                }
                metrics.update(fairness_fast(y_true, y_pred, sensitive))

                rows.append(
                    {
                        "dataset": dataset,
                        "base_model": model,
                        "seed": int(seed),
                        "rule": "Top-k per group",
                        "param": float(k),
                        "param_type": "top-k ratio per group",
                        **metrics,
                    }
                )

    candidates = pd.DataFrame(rows)
    candidates.to_csv(
        D9_DIR / "d9_postprocessing_candidates.csv",
        index=False,
    )

    mean_df = (
        candidates
        .groupby(
            [
                "dataset",
                "base_model",
                "rule",
                "param",
                "param_type",
            ],
            as_index=False,
        )
        .agg(
            pr_auc_mean=("pr_auc", "mean"),
            pr_auc_std=("pr_auc", "std"),
            f1_mean=("f1", "mean"),
            f1_std=("f1", "std"),
            eo_gap_mean=("eo_gap", "mean"),
            eo_gap_std=("eo_gap", "std"),
            delta_fpr_mean=("delta_fpr", "mean"),
            delta_fpr_std=("delta_fpr", "std"),
            delta_fnr_mean=("delta_fnr", "mean"),
            delta_fnr_std=("delta_fnr", "std"),
        )
    )

    mean_df["is_pareto"] = False

    for dataset, idx in mean_df.groupby("dataset").groups.items():
        sub = mean_df.loc[list(idx)]
        mean_df.loc[sub.index, "is_pareto"] = is_pareto_optimal(sub).values

    mean_df = add_tradeoff_scores(mean_df, base_models)

    # Chọn best trade-off theo từng rule
    best_rows = []

    for _, sub in mean_df.groupby(["dataset", "rule"]):
        guarded = sub[sub["passes_f1_guardrail"]]
        source = guarded if not guarded.empty else sub
        best_rows.append(
            source.loc[source["distance_to_ideal"].idxmin()].copy()
        )

    best_per_rule = (
        pd.DataFrame(best_rows)
        .sort_values(["dataset", "rule"])
        .reset_index(drop=True)
    )

    mean_df["is_selected_per_rule"] = False

    selected_keys = set(
        tuple(x)
        for x in best_per_rule[
            ["dataset", "base_model", "rule", "param"]
        ].to_numpy()
    )

    mean_df["is_selected_per_rule"] = [
        tuple(x) in selected_keys
        for x in mean_df[
            ["dataset", "base_model", "rule", "param"]
        ].to_numpy()
    ]

    # Chọn nghiệm khuyến nghị chung cho từng dataset
    recommended_rows = []

    for dataset, sub in best_per_rule.groupby("dataset"):
        chosen = sub.loc[sub["distance_to_ideal"].idxmin()].copy()
        recommended_rows.append(chosen)

    recommended = pd.DataFrame(recommended_rows)

    recommended_keys = set(
        tuple(x)
        for x in recommended[
            ["dataset", "base_model", "rule", "param"]
        ].to_numpy()
    )

    mean_df["is_overall_recommended"] = [
        tuple(x) in recommended_keys
        for x in mean_df[
            ["dataset", "base_model", "rule", "param"]
        ].to_numpy()
    ]

    mean_df.to_csv(
        D9_DIR / "d9_pareto_mean.csv",
        index=False,
    )

    best_per_rule.to_csv(
        D9_DIR / "D9_best_tradeoff_table.csv",
        index=False,
    )

    # So sánh baseline và post-processing khuyến nghị
    compare_rows = []

    for _, row in recommended.iterrows():
        dataset = row["dataset"]

        base = base_models[base_models["dataset"] == dataset].iloc[0]

        compare_rows.append(
            {
                "dataset": dataset,
                "base_model": row["base_model"],
                "baseline_pr_auc": base["mean_pr_auc"],
                "baseline_f1": base["mean_f1"],
                "baseline_eo_gap": base["mean_eo_gap"],
                "recommended_rule": row["rule"],
                "recommended_param": row["param"],
                "post_pr_auc": row["pr_auc_mean"],
                "post_f1": row["f1_mean"],
                "post_eo_gap": row["eo_gap_mean"],
                "eo_gap_reduction": (
                    base["mean_eo_gap"] - row["eo_gap_mean"]
                ),
                "pr_auc_change": (
                    row["pr_auc_mean"] - base["mean_pr_auc"]
                ),
            }
        )

    compare = pd.DataFrame(compare_rows)

    compare.to_csv(
        D9_DIR / "D9_baseline_vs_recommended_postprocessing.csv",
        index=False,
    )

    # Vẽ hình Pareto cho từng dataset
    for dataset in sorted(mean_df["dataset"].unique()):
        plot_dataset(
            mean_df,
            dataset,
            D9_DIR / f"pareto_{dataset}.png",
        )

    # Viết báo cáo Markdown
    report_md = D9_DIR / "D9_Pareto_Front_Report.md"
    write_report_md(best_per_rule, compare, report_md)

    # Copy báo cáo ra thư mục gốc cho dễ mở
    root_report_md = ROOT / "D9_Pareto_Front_Report.md"
    root_report_md.write_text(
        report_md.read_text(encoding="utf-8"),
        encoding="utf-8",
    )

    # Xuất thêm Word nếu có python-docx
    try_write_docx(
        report_md,
        D9_DIR / "D9_Pareto_Front_Report.docx",
    )
    try_write_docx(
        root_report_md,
        ROOT / "D9_Pareto_Front_Report.docx",
    )

    print("[OK] Đã tạo xong D9.")
    print(f"[OK] Kết quả nằm tại: {D9_DIR}")
    print("[OK] File chính:")
    print(" - results/d9/D9_best_tradeoff_table.csv")
    print(" - results/d9/D9_baseline_vs_recommended_postprocessing.csv")
    print(" - results/d9/pareto_adult.png")
    print(" - results/d9/pareto_credit_default.png")
    print(" - results/d9/D9_Pareto_Front_Report.md")
    print(" - results/d9/D9_Pareto_Front_Report.docx nếu đã cài python-docx")


if __name__ == "__main__":
    main()